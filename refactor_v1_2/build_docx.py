"""Monta o TCC v1.3 .docx (e PDF irmao) aplicando todas as refatorações.

Regras de formato (SiBiUFAL / ABNT NBR 14.724:2024):
  - Papel A4
  - Margens 3 cm sup/esq, 2 cm inf/dir
  - Fonte Arial 12 (10 para citações longas, notas, legendas, tabelas)
  - Espacamento 1,5 (simples para resumo, refs, legendas)
  - Pre/pós-textuais centralizados; seções a esquerda
  - Numeração arabica a partir da introdução
  - NUNCA usar em-dash (--) no texto

Saída:
  ../TCC - MF e G - v1.3.docx
  ../TCC - MF e G - v1.3.pdf  (via LibreOffice headless)
"""
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent
PARENT = ROOT.parent
OUT = PARENT / "TCC - MF e G - v2.7.docx"
FIG = ROOT / "figures"


# --------- helpers de formato ---------

def set_margins(section):
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)


def set_pf(p, *, align=None, line_spacing=1.5,
           space_after=Pt(0), space_before=Pt(0),
           first_line=None, left_indent=None, right_indent=None):
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.line_spacing = line_spacing
    pf.space_after = space_after
    pf.space_before = space_before
    if first_line is not None:
        pf.first_line_indent = first_line
    if left_indent is not None:
        pf.left_indent = left_indent
    if right_indent is not None:
        pf.right_indent = right_indent


def add_text(doc, text, *,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             font_size=12, bold=False, italic=False,
             line_spacing=1.5, first_line=Cm(1.25),
             space_after=Pt(6)):
    p = doc.add_paragraph()
    set_pf(p, align=align, line_spacing=line_spacing,
           space_after=space_after, first_line=first_line)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    set_pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
           space_before=Pt(12), space_after=Pt(12), first_line=Cm(0))
    run = p.add_run(text.upper())
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True


def add_h2(doc, text):
    p = doc.add_paragraph()
    set_pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5,
           space_before=Pt(12), space_after=Pt(6), first_line=Cm(0))
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True


def add_centered(doc, text, *, bold=False, size=12, line_spacing=1.5,
                 space_after=Pt(6)):
    p = doc.add_paragraph()
    set_pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=line_spacing,
           space_after=space_after, first_line=Cm(0))
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold


def add_pagebreak(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_figure(doc, path, caption, source):
    if not path.exists():
        print(f"[!] Imagem ausente: {path}")
        return
    pc = doc.add_paragraph()
    set_pf(pc, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0,
           space_after=Pt(3), first_line=Cm(0))
    r = pc.add_run(caption)
    r.font.name = "Arial"
    r.font.size = Pt(10)

    pi = doc.add_paragraph()
    set_pf(pi, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0,
           space_after=Pt(3), first_line=Cm(0))
    pi.add_run().add_picture(str(path), width=Cm(15))

    ps = doc.add_paragraph()
    set_pf(ps, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0,
           space_after=Pt(12), first_line=Cm(0))
    r = ps.add_run(source)
    r.font.name = "Arial"
    r.font.size = Pt(10)


def add_table(doc, header, rows, caption, source):
    pc = doc.add_paragraph()
    set_pf(pc, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0,
           space_after=Pt(3), first_line=Cm(0))
    r = pc.add_run(caption)
    r.font.name = "Arial"
    r.font.size = Pt(10)

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.add_run(h)
        rn.font.name = "Arial"
        rn.font.size = Pt(10)
        rn.bold = True
    for ridx, row in enumerate(rows, start=1):
        for cidx, val in enumerate(row):
            cell = table.rows[ridx].cells[cidx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            rn = p.add_run(str(val))
            rn.font.name = "Arial"
            rn.font.size = Pt(10)

    ps = doc.add_paragraph()
    set_pf(ps, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0,
           space_after=Pt(12), first_line=Cm(0))
    r = ps.add_run(source)
    r.font.name = "Arial"
    r.font.size = Pt(10)


# --------- conteúdo do TCC ---------

TITLE = (
    "INTERPRETAÇÃO DE QUESTIONÁRIOS PSICOSSOCIAIS COM MODELOS DE "
    "LINGUAGEM DE GRANDE ESCALA: UMA ARQUITETURA DE AGENTES COM RAG "
    "E VALIDAÇÃO ANTI-PII PARA A NR-01"
)


def build_pretextuais(doc):
    # Capa
    add_centered(doc, "CENTRO UNIVERSITÁRIO AFYA UNIMA-AL", bold=True)
    add_centered(doc, "CURSO DE CIÊNCIA DA COMPUTAÇÃO", bold=True)
    # Logo Afya UNIMA - AL
    logo = FIG / "logo_cesmac.png"
    if logo.exists():
        for _ in range(2):
            doc.add_paragraph()
        p = doc.add_paragraph()
        set_pf(p, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.0,
               space_after=Pt(6), first_line=Cm(0))
        p.add_run().add_picture(str(logo), width=Cm(8))
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "GUILHERME ROMUALDO MOREIRA DE ALENCAR", bold=True)
    add_centered(doc, "MARIA FERNANDA JATOBÁ CABRAL DE OLIVEIRA", bold=True)
    for _ in range(5):
        doc.add_paragraph()
    add_centered(doc,
                 "INTERPRETAÇÃO DE QUESTIONÁRIOS PSICOSSOCIAIS "
                 "COM MODELOS DE LINGUAGEM DE GRANDE ESCALA: "
                 "UMA ARQUITETURA DE AGENTES COM RAG E "
                 "VALIDAÇÃO ANTI-PII PARA A NR-01",
                 bold=True, size=14)
    for _ in range(8):
        doc.add_paragraph()
    add_centered(doc, "MACEIÓ, JUNHO DE 2026.", bold=True)
    add_pagebreak(doc)

    # Folha de rosto
    add_centered(doc, "GUILHERME ROMUALDO MOREIRA DE ALENCAR", bold=True)
    add_centered(doc, "MARIA FERNANDA JATOBÁ CABRAL DE OLIVEIRA", bold=True)
    for _ in range(6):
        doc.add_paragraph()
    add_centered(doc, TITLE, bold=True, size=14)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    set_pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0,
           space_after=Pt(0), left_indent=Cm(8), first_line=Cm(0))
    r = p.add_run(
        "Trabalho de Conclusão de Curso apresentado ao Centro "
        "Universitário Afya UNIMA-AL como um dos pré-requisitos "
        "para a obtenção do grau de Bacharel em Ciência da "
        "Computação."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0,
           space_after=Pt(0), left_indent=Cm(8), first_line=Cm(0))
    r = p.add_run("Orientador: Prof. Marcos Vinicius Silva Bento, Me.")
    r.font.name = "Arial"
    r.font.size = Pt(10)
    for _ in range(6):
        doc.add_paragraph()
    add_centered(doc, "Maceió, AL")
    add_centered(doc, "2026")
    add_pagebreak(doc)

    # Folha de aprovação
    add_centered(doc, "GUILHERME ROMUALDO MOREIRA DE ALENCAR", bold=True)
    add_centered(doc, "MARIA FERNANDA JATOBÁ CABRAL DE OLIVEIRA", bold=True)
    doc.add_paragraph()
    add_centered(doc, TITLE, bold=True, size=12)
    doc.add_paragraph()
    p = doc.add_paragraph()
    set_pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0,
           space_after=Pt(0), left_indent=Cm(8), first_line=Cm(0))
    r = p.add_run(
        "Monografia apresentada ao Centro Universitário Afya UNIMA-AL "
        "como um dos pré-requisitos para a obtenção do grau de "
        "bacharel em Ciência da Computação."
    )
    r.font.name = "Arial"
    r.font.size = Pt(10)
    for _ in range(2):
        doc.add_paragraph()
    add_centered(doc, "Aprovada em _____/_____/_____.")
    doc.add_paragraph()
    add_centered(doc, "Banca Examinadora", bold=True)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc, "_" * 60)
    add_centered(doc, "Prof. Marcos Vinicius Silva Bento, Me. (Orientador)")
    add_centered(doc, "Centro Universitário Afya UNIMA-AL")
    doc.add_paragraph()
    add_centered(doc, "_" * 60)
    add_centered(doc, "Membro Interno com Titulação")
    add_centered(doc, "Centro Universitário Afya UNIMA-AL")
    doc.add_paragraph()
    add_centered(doc, "_" * 60)
    add_centered(doc, "Membro Externo com Titulação")
    add_centered(doc, "Instituição de Ensino ou Empresa")
    add_pagebreak(doc)

    # Dedicatória e Agradecimentos
    add_h1(doc, "DEDICATÓRIA")
    add_text(doc,
             "Dedicado a todos que apoiaram esta jornada.",
             italic=True, first_line=Cm(0),
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_pagebreak(doc)
    add_h1(doc, "AGRADECIMENTOS")
    add_text(doc,
             "Aos nossos familiares, pelo suporte incondicional ao longo "
             "desta caminhada. Ao orientador, pela dedicação e pelas "
             "orientações precisas que moldaram este trabalho. Aos "
             "colegas de curso e professores do Centro Universitário "
             "Afya UNIMA-AL, pelas discussões e contribuições que "
             "enriqueceram "
             "esta pesquisa.",
             first_line=Cm(0))
    add_pagebreak(doc)

    # Resumo
    add_h1(doc, "RESUMO")
    add_text(
        doc,
        "Este trabalho investiga o uso de Modelos de Linguagem de "
        "Grande Escala (Large Language Models, LLM) com Geração "
        "Aumentada por Recuperação (Retrieval-Augmented Generation, "
        "RAG) como mecanismo de interpretação de respostas de "
        "questionários psicométricos validados de riscos psicossociais "
        "no trabalho. O problema investigado é a interpretação "
        "automatizada e auditável das respostas do Copenhagen "
        "Psychosocial Questionnaire (COPSOQ II-Br), adaptado para o "
        "português brasileiro por Boratti, Rocha e Santos (2018), "
        "visando a produção do Inventário de Riscos Psicossociais e "
        "do Plano de Ação 5W2H exigidos pelo Programa de Gerenciamento "
        "de Riscos da Norma Regulamentadora NR-01 atualizada em 2022. "
        "Foram concebidos, implementados e avaliados cinco agentes "
        "especializados sobre LLM: narrador do Inventário de Riscos; "
        "gerador de Plano de Ação 5W2H ancorado na hierarquia de "
        "controle da NR-01 e nos pressupostos andragógicos de Knowles; "
        "copiloto conversacional NR-01 com RAG sobre normas técnicas "
        "vetorizadas; validador anti-PII pós-geração baseado em "
        "expressões regulares e heurísticas semânticas, que garante "
        "anonimato dos respondentes em conformidade com a Lei Geral "
        "de Proteção de Dados (LGPD); e analisador qualitativo de "
        "respostas abertas por agrupamento semântico temático. A "
        "avaliação foi conduzida integralmente sobre dados sintéticos "
        "algoritmicamente gerados, configurando estudo de viabilidade "
        "técnica preliminar; nenhum dado real de respondente humano foi "
        "utilizado, e a validação com dados de campo permanece como "
        "trabalho futuro sujeito a aprovação por comitê de ética. O "
        "conjunto inclui 130 colaboradores fictícios distribuídos em "
        "8 setores e 2 campanhas (190 respostas completas e 15.606 "
        "respostas individuais), além de 1.025 amostras textuais com "
        "anotação span-level para o filtro anti-PII e 80 cenários com "
        "descrição cega para o gerador 5W2H. As métricas primárias "
        "seguem o princípio de adequação à função de utilidade do "
        "problema: Fβ=2 para o validador anti-PII (em que falso negativo "
        "é catastrófico em LGPD), BERTScore F1 e taxa de alucinação "
        "regulatória para os agentes generativos, e framework RAGAS "
        "para o copiloto. Todas as métricas pontuais são acompanhadas "
        "de intervalos de confiança de 95% obtidos por bootstrap "
        "(Efron e Tibshirani, 1993, n=1.000 reamostragens). O validador "
        "anti-PII alcançou Fβ=2 de 0,9710 (IC95% [0,9563; 0,9843]) "
        "no nível de documento e Fβ=2 macro span-level de 0,9688 "
        "(IC95% [0,9546; 0,9822]), com taxa de vazamento por "
        "categoria abaixo de 0,15 e latência p99 inferior a 0,1 ms. "
        "O copiloto "
        "RAG, sob o framework RAGAS, apresentou recall@5 de 0,8250, "
        "Precision@1 de 0,6500, faithfulness de 0,1514 (IC95% [0,09; "
        "0,22]) e taxa de alucinação regulatória de 0,0000 sobre as "
        "citações produzidas, com corpus integral indexado de NR-01, "
        "NR-17 e LGPD (596 chunks). A faithfulness modesta indica "
        "que o componente RAG não ancora efetivamente as respostas "
        "nos chunks recuperados, limitação central do estudo. O "
        "trabalho inclui ainda benchmark tripartite com "
        "dois LLMs open-source de 14 bilhões de parâmetros "
        "executados localmente (Qwen 2.5 14B e Phi-4 14B via "
        "Ollama em GPU AMD), em que o Qwen supera o gpt-4o-mini "
        "no copiloto (Precision@1 = 0,8750 vs. 0,6500) e o Phi-4 "
        "atinge alucinação regulatória de zero com word range "
        "estrutural de 0,9400 no narrador, evidenciando "
        "viabilidade de migração parcial para modelos "
        "open-source. Os resultados preliminares sobre "
        "dados sintéticos sugerem viabilidade técnica da "
        "arquitetura combinada (LLM + RAG + k-anonymity + anti-PII) "
        "condicionada à correção da ancoragem do RAG e à validação "
        "com dados reais sob aprovação ética. Código-fonte e "
        "artefatos disponíveis em "
        "https://github.com/Alencar-png/emotion-care-tcc.",
        line_spacing=1.0, first_line=Cm(0))
    add_text(
        doc,
        "Palavras-chave: Large Language Models; Geração Aumentada por "
        "Recuperação; riscos psicossociais; COPSOQ-II; Norma "
        "Regulamentadora NR-01; saúde mental no trabalho; "
        "anonimização de dados.",
        line_spacing=1.0, first_line=Cm(0))
    add_pagebreak(doc)

    # Abstract
    add_h1(doc, "ABSTRACT")
    add_text(
        doc,
        "This work investigates the use of Large Language Models "
        "(LLM) with Retrieval-Augmented Generation (RAG) as an "
        "interpretation mechanism for responses to validated "
        "psychometric questionnaires on workplace psychosocial risks. "
        "The investigated problem is the automated and auditable "
        "interpretation of responses to the Copenhagen Psychosocial "
        "Questionnaire (COPSOQ II-Br), adapted to Brazilian "
        "Portuguese by Boratti, Rocha and Santos (2018), aiming at "
        "the production of the Psychosocial Risk Inventory and of "
        "the 5W2H Action Plan required by the Occupational Risk "
        "Management Program of Brazilian Regulatory Standard NR-01 "
        "(2022 revision). Five specialized LLM-based agents were "
        "designed, implemented and evaluated: a Risk Inventory "
        "narrator; a 5W2H Action Plan generator anchored to NR-01 "
        "hierarchy of controls and to Knowles andragogical principles; "
        "a conversational NR-01 copilot with RAG over vectorized "
        "regulatory documents; a post-generation anti-PII validator; "
        "and a qualitative analyzer that semantically clusters "
        "open-ended responses into themes. The evaluation was "
        "conducted entirely on algorithmically generated synthetic "
        "data, characterizing a preliminary technical feasibility "
        "study; no real respondent data was used, and field "
        "validation remains as future work subject to ethics committee "
        "approval. The dataset includes 130 fictitious employees "
        "across 8 departments and 2 campaigns (190 complete responses "
        "and 15,606 individual answers), 1,025 span-level annotated "
        "textual samples for the anti-PII validator and 80 blind-"
        "description scenarios for the 5W2H generator. Primary "
        "metrics were selected to fit each problem's utility function: "
        "Fβ=2 for the anti-PII validator (where false negatives are "
        "catastrophic under LGPD), BERTScore F1 and regulatory "
        "hallucination rate for generative agents, and the RAGAS "
        "framework for the copilot. All point metrics are accompanied "
        "by 95% confidence intervals from non-parametric bootstrap "
        "(n=1,000 resamples). The anti-PII validator reached document-"
        "level Fβ=2 = 0.9710 (95% CI [0.9563, 0.9843]) and macro span-"
        "level Fβ=2 = 0.9688 (95% CI [0.9546, 0.9822]), with category-"
        "wise leakage rate below 0.15 and p99 latency under 0.1 ms. "
        "The copilot RAG, under the RAGAS framework, reached "
        "recall@5 = 0.8250, Precision@1 = 0.6500, faithfulness = "
        "0.1514 (95% CI [0.09, 0.22]) and regulatory hallucination "
        "rate of 0.0000 over produced citations, with the integral "
        "indexing of NR-01, NR-17 and LGPD (596 chunks). Modest "
        "faithfulness indicates that the RAG component did not "
        "effectively ground answers in retrieved chunks. This "
        "work also includes a triple-model benchmark "
        "against two open-source 14B-parameter LLMs running "
        "locally (Qwen 2.5 14B and Phi-4 14B via Ollama on AMD "
        "GPU), revealing that Qwen surpasses gpt-4o-mini on the "
        "copilot (Precision@1 = 0.8750 vs. 0.6500) and Phi-4 "
        "reaches zero regulatory hallucination with structural "
        "word range of 0.9400 in the narrator, evidencing the "
        "feasibility of partial migration to open-source models. "
        "Preliminary results over synthetic data suggest the technical "
        "viability of the combined architecture (LLM + RAG + "
        "k-anonymity + anti-PII) and point to the next stage: "
        "validation with real data under ethical approval. Source "
        "code and artifacts are publicly available at "
        "https://github.com/Alencar-png/emotion-care-tcc.",
        line_spacing=1.0, first_line=Cm(0))
    add_text(
        doc,
        "Keywords: Large Language Models; Retrieval-Augmented "
        "Generation; psychosocial risks; COPSOQ-II; Brazilian "
        "Regulatory Standard NR-01; workplace mental health; data "
        "anonymization.",
        line_spacing=1.0, first_line=Cm(0))
    add_pagebreak(doc)

    # Listas
    add_h1(doc, "LISTA DE FIGURAS")
    figs = [
        "Figura 1 - System Design: arquitetura conceitual do Emotion Care",
        "Figura 2 - User Flow: colaborador respondente e gestor de SST",
        "Figura 3 - Pipeline LLM/RAG dos 5 agentes com validador anti-PII",
        "Figura 4 - Processamento de uma resposta no agente narrador do PGR",
        "Figura 5 - Matriz de confusão doc-level do validador anti-PII (n=1.025)",
        "Figura 6 - Métricas span-level por categoria do validador anti-PII",
        "Figura 7 - Matriz de confusão 4x4 da hierarquia de controle NR-01",
        "Figura 8 - Gerador 5W2H, F1 por nível NR-01 (3 modelos)",
        "Figura 9 - Copiloto NR-01, métricas RAGAS (3 modelos)",
        "Figura 10 - Narrador do PGR, métricas primárias (3 modelos)",
        "Figura 11 - Analisador qualitativo, ARI/NMI/c_v/F1 (3 modelos, 5 seeds)",
        "Figura 12 - Pareto latência mediana versus qualidade (3 modelos)",
        "Figura 13 - Anti-PII Fβ=2 por categoria, evolução do validador determinístico",
    ]
    for f in figs:
        add_text(doc, f, line_spacing=1.0, first_line=Cm(0))
    add_pagebreak(doc)

    add_h1(doc, "LISTA DE TABELAS")
    tabs = [
        "Tabela 1 - Síntese dos trabalhos relacionados e lacunas",
        "Tabela 2 - Tecnologias adotadas e respectivas versões",
        "Tabela 3 - Gold standards construídos para a avaliação dos agentes",
        "Tabela 4 - Resultados do validador anti-PII no nível de documento",
        "Tabela 5 - Métricas por categoria do validador anti-PII",
        "Tabela 5b - Taxa de vazamento por categoria do validador anti-PII",
        "Tabela 6 - Conformidade estrutural do narrador do PGR",
        "Tabela 7 - Resultados do gerador 5W2H",
        "Tabela 7b - Métricas por nível NR-01 do gerador 5W2H",
        "Tabela 8 - Recuperação do copiloto NR-01",
        "Tabela 9 - Métricas globais do analisador qualitativo",
        "Tabela 9b - Desempenho por tema do analisador qualitativo",
        "Tabela 10 - Métricas operacionais por agente",
        "Tabela 11 - Métricas operacionais do site",
        "Tabela 12 - Comparação dos quatro agentes generativos entre gpt-4o-mini e LLMs open-source 14B",
        "Tabela 13 - Comparação descritiva do validador anti-PII com trabalhos correlatos",
    ]
    quadros = [
        "Quadro 1 - Aparato metodológico: métricas primárias por agente",
        "Quadro 2 - Módulos de software para reprodutibilidade",
    ]
    for t in tabs:
        add_text(doc, t, line_spacing=1.0, first_line=Cm(0))
    add_pagebreak(doc)

    add_h1(doc, "LISTA DE QUADROS")
    for q in quadros:
        add_text(doc, q, line_spacing=1.0, first_line=Cm(0))
    add_pagebreak(doc)

    add_h1(doc, "LISTA DE SIGLAS")
    siglas = [
        ("ABNT", "Associação Brasileira de Normas Técnicas"),
        ("API", "Application Programming Interface"),
        ("ARI", "Adjusted Rand Index"),
        ("AUC", "Area Under the Curve"),
        ("BERT", "Bidirectional Encoder Representations from Transformers"),
        ("CID", "Classificação Internacional de Doenças"),
        ("CIPA", "Comissão Interna de Prevenção de Acidentes"),
        ("CNPJ", "Cadastro Nacional da Pessoa Jurídica"),
        ("COPSOQ", "Copenhagen Psychosocial Questionnaire"),
        ("CPF", "Cadastro de Pessoas Físicas"),
        ("DASS", "Depression Anxiety Stress Scales"),
        ("EACT", "Escala de Avaliação do Contexto do Trabalho"),
        ("EET", "Escala de Estresse no Trabalho"),
        ("GHE", "Grupo Homogêneo de Exposição"),
        ("GRO", "Gerenciamento de Riscos Ocupacionais"),
        ("INSS", "Instituto Nacional do Seguro Social"),
        ("ISO", "International Organization for Standardization"),
        ("LDA", "Latent Dirichlet Allocation"),
        ("LGPD", "Lei Geral de Proteção de Dados Pessoais"),
        ("LLM", "Large Language Model"),
        ("MBI", "Maslach Burnout Inventory"),
        ("MRR", "Mean Reciprocal Rank"),
        ("NER", "Named Entity Recognition"),
        ("NMI", "Normalized Mutual Information"),
        ("NR", "Norma Regulamentadora"),
        ("OMS", "Organização Mundial da Saúde"),
        ("OIT", "Organização Internacional do Trabalho"),
        ("PGR", "Programa de Gerenciamento de Riscos"),
        ("PHQ", "Patient Health Questionnaire"),
        ("PII", "Personally Identifiable Information"),
        ("QPS", "Questionnaire on Psychological and Social Factors at Work"),
        ("RAG", "Retrieval-Augmented Generation"),
        ("REST", "Representational State Transfer"),
        ("ROC", "Receiver Operating Characteristic"),
        ("RNN", "Recurrent Neural Network"),
        ("SaaS", "Software as a Service"),
        ("SST", "Saúde e Segurança do Trabalho"),
        ("TF-IDF", "Term Frequency-Inverse Document Frequency"),
    ]
    for s, n in siglas:
        add_text(doc, f"{s}: {n}", line_spacing=1.0, first_line=Cm(0))
    add_pagebreak(doc)

    add_h1(doc, "SUMÁRIO")
    toc = [
        "1 INTRODUÇÃO",
        "2 OBJETIVOS",
        "2.1 Objetivo geral",
        "2.2 Objetivos específicos",
        "3 ASPECTOS TEÓRICOS",
        "3.1 Saúde mental no trabalho e riscos psicossociais",
        "3.2 Base regulatória e instrumentos de avaliação",
        "3.3 Modelos Transformer e RAG",
        "3.4 Indicadores, dashboards e apoio a decisão",
        "3.5 Privacidade e k-anonymity",
        "4 TRABALHOS RELACIONADOS",
        "4.1 ML aplicado a instrumentos psicométricos",
        "4.2 RAG em domínios regulados",
        "4.3 Geração automatizada de relatórios técnicos",
        "4.4 Anonimização e validação anti-PII em LLM",
        "4.5 Análise temática de respostas abertas",
        "4.6 Síntese e lacunas identificadas",
        "5 METODOLOGIA",
        "5.1 Ambiente de desenvolvimento",
        "5.2 Tecnologias adotadas",
        "5.3 Configurações do projeto",
        "5.4 Fluxo da solução",
        "5.5 Construção dos gold standards",
        "5.6 Métricas adotadas",
        "5.7 Garantias de segurança e privacidade do experimento",
        "6 RESULTADOS",
        "6.1 Validador anti-PII",
        "6.2 Narrador do PGR",
        "6.3 Gerador de Plano de Ação 5W2H",
        "6.4 Copiloto NR-01 com RAG",
        "6.5 Analisador qualitativo",
        "6.6 Métricas operacionais do sistema",
        "6.7 Métricas de uso da plataforma",
        "7 DISCUSSÃO",
        "7.1 Anonimização em comparação com trabalhos correlatos",
        "7.2 Geração de documentos técnicos",
        "7.3 Aderência à hierarquia de controle NR-01",
        "7.4 RAG sobre normas brasileiras vs. MIRAGE",
        "7.5 Clustering temático",
        "7.6 Limitações observadas e ameaças à validade",
        "7.7 Implicações práticas e científicas",
        "8 CONCLUSÕES",
        "9 SUGESTÕES PARA TRABALHOS FUTUROS",
        "9.1 Validação com dados reais sob aprovação ética",
        "9.2 Reranking pós-retrieval e fine-tuning de ancoragem",
        "9.3 Fine-tuning para o Nível 2 da hierarquia 5W2H",
        "9.4 Aplicação Likert humana e medição de kappa",
        "9.5 Extensões de escopo",
        "REFERÊNCIAS",
        "APÊNDICE A - MÓDULOS DE SOFTWARE PARA REPRODUTIBILIDADE",
        "APÊNDICE B - PROMPTS DE SISTEMA DOS AGENTES LLM",
    ]
    for entry in toc:
        add_text(doc, entry, line_spacing=1.0, first_line=Cm(0))
    add_pagebreak(doc)


def build_introducao(doc):
    add_h1(doc, "1 INTRODUÇÃO")
    paras = [
        "A saúde mental no trabalho passou a ocupar papel estratégico nas organizações em razão do aumento expressivo dos afastamentos relacionados a transtornos mentais e comportamentais. Segundo o Anuário Estatístico Previdenciário do Instituto Nacional do Seguro Social (INSS, 2024), os transtornos mentais e comportamentais (Capítulo V da CID-10, classes F00 a F99) figuraram entre as três principais causas de afastamento do trabalho no Brasil, com mais de 200 mil benefícios concedidos no período. Esse cenário se agravou no período pós-pandemia e levou organizações internacionais como a Organização Mundial da Saúde (WHO, 2022) e a Organização Internacional do Trabalho (ILO, 2022) a publicar diretrizes específicas para a proteção da saúde mental no ambiente laboral.",
        "No Brasil, a atualização da Norma Regulamentadora NR-01, publicada pelo Ministério do Trabalho e Emprego em 2022, instituiu o Gerenciamento de Riscos Ocupacionais (GRO). Posteriormente, a Portaria MTE nº 1.419/2024 incluiu de forma explícita os fatores de risco psicossocial entre os perigos a serem identificados, avaliados e controlados no Programa de Gerenciamento de Riscos (PGR). A Portaria MTE nº 765/2025 fixou a vigência fiscalizável dessa obrigação em 26 de maio de 2026, mantendo o período anterior em caráter educativo. Soma-se a esse arcabouço a NR-17 (BRASIL, 2018), que trata dos fatores ergonômicos cognitivos e organizacionais, e a norma internacional ISO 45003 (2021), dedicada exclusivamente a saúde e segurança psicológica no trabalho. Nesse contexto, e considerando que o presente trabalho é entregue pouco depois do início da exigibilidade da obrigação, empresas, clínicas de medicina ocupacional e consultorias de Saúde e Segurança do Trabalho (SST) precisam realizar avaliações periódicas com instrumentos psicométricos validados, interpretar os resultados, gerar evidências auditáveis de conformidade e transformar os achados em planos de ação efetivos. Entretanto, a interpretação das respostas e a redação dos documentos exigidos pela NR-01 permanecem majoritariamente manuais, demoradas, dependentes da subjetividade de cada profissional de SST e com baixa rastreabilidade do raciocínio empregado.",
        "A maturação dos Modelos de Linguagem de Grande Escala (Large Language Models, LLM) baseados em arquiteturas Transformer (VASWANI et al., 2017) e, em especial, da técnica de Geração Aumentada por Recuperação (Retrieval-Augmented Generation, RAG), proposta por Lewis et al. (2020), abriu novas possibilidades para a automação da interpretação de instrumentos psicométricos. Em vez de uma simples mecanização de fórmulas, as LLM permitem traduzir matrizes de escores em narrativas técnicas auditáveis, ancorar respostas em fontes normativas recuperadas dinamicamente e agrupar respostas abertas por similaridade semântica sem expor respondentes individualmente. No campo da saúde ocupacional, todavia, sua adoção esbarra em três restrições não negociáveis: a hipersensibilidade dos dados (LGPD trata dados de saúde como sensíveis), o risco de alucinação característico das LLM e a exigência de que cada conclusão seja rastreável até evidências aceitas pelo arcabouço normativo.",
        "Este trabalho delimita-se a um estudo de viabilidade técnica conduzido integralmente sobre dados sintéticos algoritmicamente gerados. Nenhum respondente humano foi avaliado e nenhum dado clínico ou ocupacional real foi coletado; a validação em campo, sujeita a aprovação por comitê de ética em pesquisa, permanece como trabalho futuro inegociável.",
    ]
    for txt in paras:
        add_text(doc, txt)

    # Pergunta-problema em destaque
    p = doc.add_paragraph()
    set_pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
           space_after=Pt(12), first_line=Cm(0),
           left_indent=Cm(2), right_indent=Cm(2))
    r1 = p.add_run("Diante desse cenário, formula-se a seguinte pergunta-problema: ")
    r1.font.name = "Arial"
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(
        "é possível utilizar Modelos de Linguagem de Grande Escala, "
        "combinados a Geração Aumentada por Recuperação e a um piso "
        "inegociável de k-anonymity, para interpretar respostas do "
        "COPSOQ II-Br e produzir automaticamente o Inventário de "
        "Riscos Psicossociais e o Plano de Ação 5W2H exigidos pela "
        "NR-01, preservando o anonimato dos respondentes e a "
        "rastreabilidade documental requerida em auditorias "
        "trabalhistas?"
    )
    r2.font.name = "Arial"
    r2.font.size = Pt(12)
    r2.italic = True

    paras2 = [
        "A presente pesquisa endereça essa pergunta com a concepção, a implementação e a avaliação quantitativa de cinco agentes especializados sobre LLM: um narrador do Inventário de Riscos Psicossociais, um gerador de Plano de Ação 5W2H ancorado na hierarquia de controle da NR-01, um copiloto conversacional NR-01 baseado em RAG sobre normas técnicas vetorizadas, um validador anti-PII que opera como filtro pós-geração obrigatório sobre cada saída dos demais agentes, e um analisador qualitativo que agrupa respostas abertas por similaridade semântica em temas anônimos. A plataforma Emotion Care, construída como ambiente experimental e como veículo de exposição dos agentes, é descrita na Metodologia apenas no nível necessário a reprodutibilidade do experimento, sem ocupar o centro da discussão.",
        "A relevância do estudo reside na intersecção, ainda pouco explorada na literatura brasileira, entre instrumentos psicométricos validados, arcabouço regulatório de SST e LLM com RAG em domínios sensíveis. Trabalhos correlatos (apresentados no Capítulo 4) reportam o uso de modelos clássicos de aprendizado de máquina para classificação de risco psicossocial e de LLM em domínios médicos genéricos, porém ainda são escassos os estudos que aplicam LLM diretamente ao COPSOQ-II em português, integram RAG sobre o corpus normativo brasileiro de SST e reportam métricas quantitativas de aderência das saídas geradas a um gold standard avaliado por especialistas.",
        "O presente trabalho está organizado em nove capítulos. O Capítulo 2 enuncia os objetivos geral e específicos. O Capítulo 3 reúne os aspectos teóricos. O Capítulo 4 apresenta os trabalhos relacionados, distinto dos aspectos teóricos. O Capítulo 5 descreve a metodologia. O Capítulo 6 apresenta os resultados quantitativos. O Capítulo 7 discute os resultados em comparação com os trabalhos relacionados. O Capítulo 8 consolida as conclusões. O Capítulo 9 elenca trabalhos futuros.",
    ]
    for txt in paras2:
        add_text(doc, txt)


def build_objetivos(doc):
    add_h1(doc, "2 OBJETIVOS")
    add_h2(doc, "2.1 Objetivo geral")
    add_text(
        doc,
        "Investigar, projetar e avaliar quantitativamente o uso de "
        "Modelos de Linguagem de Grande Escala (LLM) combinados a "
        "técnica de Geração Aumentada por Recuperação (RAG) como "
        "mecanismo de interpretação de respostas do questionário "
        "COPSOQ II-Br, com vistas a produção automatizada do "
        "Inventário de Riscos Psicossociais e do Plano de Ação 5W2H "
        "exigidos pelo Programa de Gerenciamento de Riscos da Norma "
        "Regulamentadora NR-01, preservando o anonimato dos "
        "respondentes em conformidade com a Lei Geral de Proteção de "
        "Dados (LGPD)."
    )
    add_h2(doc, "2.2 Objetivos específicos")
    items = [
        "a) Levantar, na literatura científica e no arcabouço normativo, requisitos para a interpretação automatizada de instrumentos psicométricos de riscos psicossociais e identificar as lacunas que justifiquem a abordagem baseada em LLM com RAG;",
        "b) Especificar cinco agentes especializados sobre LLM cobrindo as operações críticas do processo: narração do Inventário de Riscos, geração do Plano de Ação 5W2H ancorado na hierarquia de controle da NR-01, copiloto conversacional sobre normas técnicas via RAG, validação anti-PII pós-geração e agrupamento semântico de respostas abertas;",
        "c) Definir um motor analítico determinístico que normalize as respostas do COPSOQ II-Br, agregue indicadores por dimensão e por setor e imponha um piso inegociável de k-anonymity antes de qualquer exposição a usuários ou aos agentes LLM;",
        "d) Implementar o sistema experimental Emotion Care como ambiente controlado de execução, configuração e instrumentação dos agentes, em nível suficiente para reprodução do experimento e captura de métricas operacionais;",
        "e) Definir e executar um protocolo de avaliação dos agentes com base em métricas quantitativas justificadas pela função de utilidade de cada problema (precisão, recall, F1-score, Fβ=2, BERTScore, faithfulness, RAGAS e coerência tópica c_v) com intervalos de confiança 95% por bootstrap, sobre gold standards construídos para cada agente;",
        "f) Comparar os resultados obtidos com as métricas reportadas pelos trabalhos relacionados, evidenciando ganhos, limitações e oportunidades de evolução da abordagem.",
    ]
    for it in items:
        add_text(doc, it)


def build_aspectos_teoricos(doc):
    add_h1(doc, "3 ASPECTOS TEÓRICOS")
    add_h2(doc, "3.1 Saúde mental no trabalho e riscos psicossociais")
    add_text(
        doc,
        "A saúde mental no ambiente corporativo é definida pela "
        "Organização Mundial da Saúde (WHO, 2022) como um estado de "
        "bem-estar no qual o trabalhador percebe suas próprias "
        "capacidades, lida com o estresse normal da vida, é produtivo "
        "e consegue contribuir para sua comunidade. No contexto "
        "laboral, esse estado é afetado por fatores ligados a carga "
        "de trabalho, ao ritmo, as relações interpessoais, ao estilo "
        "de liderança, a organização do trabalho e as condições "
        "ergonômicas. A literatura consolidou três modelos teóricos "
        "estruturantes para o estudo do estresse ocupacional: o "
        "Modelo Demanda-Controle de Karasek e Theorell (1990); o "
        "Modelo Esforço-Recompensa de Siegrist (1996); e a abordagem "
        "de Cox e Griffiths (1995). Quando esses fatores não são "
        "monitorados, podem resultar em queda de produtividade, "
        "absenteísmo, presenteísmo, burnout (Maslach e Jackson, 1981) "
        "e afastamentos."
    )
    add_h2(doc, "3.2 Base regulatória e instrumentos de avaliação")
    add_text(
        doc,
        "A plataforma proposta tem como referência regulatória três "
        "pilares principais. O primeiro é a Norma Regulamentadora "
        "NR-01 (BRASIL, 2022), que instituiu o Gerenciamento de "
        "Riscos Ocupacionais (GRO), e a Portaria MTE nº 1.419/2024 "
        "(BRASIL, 2024), que incluiu de forma explícita os fatores "
        "de risco psicossocial entre os perigos a serem "
        "identificados, avaliados e controlados no Programa de "
        "Gerenciamento de Riscos (PGR), com vigência fiscalizável "
        "a partir de 26 de maio de 2026 fixada pela Portaria MTE "
        "nº 765/2025 (BRASIL, 2025). O segundo é a NR-17 "
        "(BRASIL, 2018), que trata dos fatores "
        "ergonômicos cognitivos e organizacionais. O terceiro é a "
        "ISO 45003:2021, dedicada a saúde e segurança psicológica no "
        "trabalho. A esses pilares soma-se a Lei Geral de Proteção "
        "de Dados Pessoais (BRASIL, 2018), que classifica dados de "
        "saúde como sensíveis. Para operacionalizar a avaliação, são "
        "amplamente reconhecidos o Copenhagen Psychosocial "
        "Questionnaire em sua segunda versão (COPSOQ-II), proposto "
        "por Kristensen et al. (2010) e adaptado por Boratti, Rocha "
        "e Santos (2018) em versão composta por 26 dimensões e 80 "
        "questões em escala Likert de cinco pontos."
    )
    add_h2(doc, "3.3 Modelos Transformer e Geração Aumentada por Recuperação")
    add_text(
        doc,
        "Os Modelos de Linguagem de Grande Escala (LLM) baseiam-se "
        "em arquiteturas Transformer (VASWANI et al., 2017), nas "
        "quais a operação central é o mecanismo de autoatenção que "
        "pondera, para cada token de entrada, sua relação com todos "
        "os demais tokens da sequência. A arquitetura é composta por "
        "blocos de atenção multi-cabeca intercalados com camadas "
        "densas e normalizações, configurando um decodificador "
        "autoregressivo. A Geração Aumentada por Recuperação "
        "(Retrieval-Augmented Generation, RAG), proposta por Lewis "
        "et al. (2020), estende o paradigma incorporando uma etapa "
        "de recuperação de documentos relevantes antes da geração: "
        "uma consulta e vetorizada via modelo de embeddings, "
        "comparada por similaridade de cosseno a um índice de "
        "vetores e os trechos mais relevantes são adicionados ao "
        "prompt. A combinação reduz alucinações e ancora a saída em "
        "fontes auditáveis, propriedade essencial em domínios "
        "regulados como saúde ocupacional."
    )
    add_h2(doc, "3.4 Indicadores, dashboards e apoio a decisão")
    add_text(
        doc,
        "A transformação das respostas dos questionários em médias, "
        "percentuais de risco e escores gerais permite construir "
        "dashboards com visao consolidada por empresa, setor, cargo "
        "e ciclo de avaliação. Tufte (2001) e Few (2009) "
        "consolidaram principios para a construção de paineis "
        "informativos eficazes: alta razão dado-tinta, supressão de "
        "elementos decorativos, enquadramento consistente do "
        "contexto temporal e uso parcimonioso de cores semânticas. "
        "O motor de análise da plataforma proposta normaliza cada "
        "resposta Likert para uma escala de zero a cem (com inversão "
        "para itens reverso), agrega valores por dimensão, "
        "departamento e cargo, e classifica cada agregado em três "
        "níveis de risco (saudável, intermediario e em risco) "
        "segundo a convenção do COPSOQ-II."
    )
    add_h2(doc, "3.5 Privacidade e k-anonymity")
    add_text(
        doc,
        "Sweeney (2002) introduziu o modelo de k-anonymity para "
        "proteção de privacidade em conjuntos de dados liberados a "
        "terceiros: uma liberação satisfaz k-anonymity se, para cada "
        "indivíduo, suas informações são indistinguíveis de pelo "
        "menos k-1 outros indivíduos no mesmo conjunto. Aplicado ao "
        "contexto desta pesquisa, o piso inviolável de k=4 (scores) "
        "e k=3 (recortes setoriais) impede que qualquer agregado "
        "venha a ser exposto se a contagem do recorte estiver abaixo "
        "desses limites, mitigando o risco de reidentificação por "
        "inferência. Combinado ao validador anti-PII pós-geração, o "
        "piso de k-anonymity constitui a base de defesa em "
        "profundidade adotada neste trabalho."
    )
    add_text(
        doc,
        "A escolha específica de k=4 para scores agregados por "
        "dimensão e k=3 para recortes setoriais reflete o trade-off "
        "entre utilidade analítica e risco de reidentificação. "
        "Valores menores (k=1 ou k=2) preservariam mais detalhe nas "
        "agregações mas permitiriam identificação por exclusão em "
        "setores pequenos, particularmente em empresas de até 50 "
        "colaboradores onde um único respondente em determinada "
        "função pode ser inferido. Valores maiores (k=10 ou k=20) "
        "ofereceriam maior privacidade ao custo de impedir análise "
        "em setores legítimos com 5 a 9 respondentes, comuns em "
        "empresas brasileiras de pequeno porte. O piso k=4 para "
        "scores e k=3 para setores foi calibrado para preservar a "
        "utilidade do Inventário em organizações com 30 a 100 "
        "colaboradores sem comprometer o anonimato dos respondentes "
        "individuais. A política é inegociável: parâmetros podem "
        "ser ajustados para baixo apenas em direção mais "
        "conservadora (k=5, k=6), nunca para baixo do piso "
        "estabelecido."
    )


def build_trabalhos_relacionados(doc):
    add_h1(doc, "4 TRABALHOS RELACIONADOS")
    add_text(
        doc,
        "A presente seção apresenta os trabalhos correlatos que "
        "serviram de referência para a concepção dos cinco agentes "
        "propostos e para a definição do protocolo de avaliação. A "
        "revisão foi conduzida nas bases PubMed, IEEE Xplore, ACM "
        "Digital Library, ACL Anthology, SciELO e arXiv, com recorte "
        "temporal entre 2018 e 2026. Os trabalhos foram selecionados "
        "por reportarem métricas quantitativas, apresentarem técnica "
        "de processamento de linguagem natural ou aprendizado de "
        "máquina aplicada a domínios relacionados a saúde mental, "
        "saúde ocupacional, normas regulatórias ou anonimização, e "
        "estarem acessíveis em texto completo."
    )

    add_h2(doc, "4.1 ML aplicado a instrumentos psicométricos")
    add_text(
        doc,
        "Shi et al. (2025), em meta-análise de 22 estudos primarios "
        "sobre predição de burnout em profissionais de saúde com o "
        "Maslach Burnout Inventory como ground truth, reportam AUC "
        "agrupada de 0,72 (IC95% 0,68 a 0,76), sensibilidade de 0,63 "
        "e especificidade de 0,84. AlSaad, Alshakhs e Thomas (2026) "
        "compararam, sobre 14.983 tweets em inglês rotulados em 6 "
        "subtipos de depressão, few-shot prompting de LLMs (Llama-3, "
        "Mistral, Phi-3.5) contra fine-tuning com LoRA sobre "
        "encoders clássicos, atingindo F1 macro de 0,957 contra "
        "0,765 das LLMs em few-shot. Shao et al. (2025) avaliaram o "
        "PHQ-9 sobre tweets decompondo a detecção em 11 subtarefas, "
        "com melhor F1 de 0,872 mas apenas 17,8 a 51,3 por cento de "
        "correção conjunta nas 11 tarefas."
    )

    add_h2(doc, "4.2 RAG em domínios regulados")
    add_text(
        doc,
        "Xiong et al. (2024) apresentaram o MIRAGE, benchmark de RAG "
        "para a medicina composto por 7.663 questões e 41 "
        "combinações de corpus, retriever e LLM, demonstrando que "
        "RAG eleva a acurácia em 18 pontos percentuais sobre "
        "Chain-of-Thought puro. Hillebrand et al. (2025) aplicaram "
        "RAG na construção de chatbots para compliance regulatório "
        "financeiro e relataram redução significativa na taxa de "
        "alucinação, reforçando a lacuna metodológica no domínio "
        "regulatório de SST brasileiro."
    )

    add_h2(doc, "4.3 Geração automatizada de relatórios técnicos")
    add_text(
        doc,
        "Voinea et al. (2024) demonstraram que Llama-3-8B fine-tuned "
        "com LoRA gera conclusões de laudos radiológicos a partir de "
        "21.152 exames, atingindo BERTScore F1 de 0,805, ROUGE-1 F1 "
        "de 0,500 e BLEU de 0,225. Em avaliação cega tipo Turing com "
        "13 radiologistas, as conclusões geradas foram preferidas em "
        "21,8 por cento dos casos sobre as escritas por humanos."
    )

    add_h2(doc, "4.4 Anonimização e validação anti-PII em LLM")
    add_text(
        doc,
        "Schiezaro et al. (2026) avaliaram pipelines de anonimização "
        "de prontuários médicos em português brasileiro comparando "
        "modelos extrativos (BERTimbau-leNER, mBERT, BioBERTpt) e "
        "generativos (ptt5-v2, GPT-4o), apoiados por framework "
        "LLM-as-a-Judge (Gemini 2.5 Pro). O melhor pipeline atingiu "
        "F1 de 0,927, precisão de 0,926 e recall de 0,9351 sobre "
        "2.962 registros. Wiegand et al. (2024), no projeto "
        "LLM-Anonymizer, propuseram de-identificação local de "
        "documentos médicos em alemão e inglês usando LLMs "
        "open-source, atingindo acurácia caractere a caractere de "
        "0,992 e recall de 0,9794."
    )

    add_h2(doc, "4.5 Análise temática de respostas abertas")
    add_text(
        doc,
        "Jiang, Liu e Fisher (2026) compararam Structural Topic "
        "Models (STM) e BERTopic sobre respostas abertas curtas em "
        "pesquisas de opinião, reportando superioridade do BERTopic "
        "em coerência tópica c_v em todas as variações testadas, "
        "com ganhos adicionais quando empregada data augmentation "
        "contextual."
    )

    add_h2(doc, "4.6 Síntese e lacunas identificadas")
    add_text(
        doc,
        "A Tabela 1 sintetiza os trabalhos analisados, suas "
        "técnicas, tamanhos de dataset e métricas reportadas, e "
        "indica qual lacuna o presente trabalho preenche."
    )

    header = ["Trabalho", "Domínio", "Técnica", "n",
              "Metrica principal", "Lacuna preenchida"]
    rows = [
        ["Shi et al. (2025)", "Burnout / MBI",
         "Meta-análise ML clássico", "22 estudos", "AUC 0,72",
         "Não interpreta semanticamente nem gera documento"],
        ["AlSaad et al. (2026)", "Subtipos de depressão",
         "Few-shot vs fine-tuning", "14.983 tweets",
         "F1 macro 0,957", "Texto livre, não instrumento validado"],
        ["Shao et al. (2025)", "PHQ-9 em redes sociais",
         "CoT/SFT/DPO em LLM 7B", "3.132 tweets", "F1 0,872",
         "Não produz saída regulatória"],
        ["Xiong et al. (2024)", "QA médico com RAG",
         "Benchmark MIRAGE", "7.663 questões",
         "Acc. +18 p.p.", "Não cobre NR-01 nem prescritivo"],
        ["Voinea et al. (2024)", "Laudo radiológico",
         "Llama-3 fine-tuned", "21.152 laudos",
         "BERTScore F1 0,805", "Outro domínio e produto final"],
        ["Schiezaro et al. (2026)", "Anonimização PT-BR clínica",
         "BERT+LLM extrativo+generativo", "2.962 registros",
         "F1 0,927", "Não integra anonimizador a gerador"],
        ["Wiegand et al. (2024)", "De-identification local",
         "LLM open-source", "~250 docs", "Recall 0,9794",
         "Não cobre geração subsequente"],
        ["Jiang et al. (2026)", "Clustering temático",
         "BERTopic vs STM", "Surveys", "Coerência c_v",
         "Não vincula clusters a psicometria"],
    ]
    add_table(doc, header, rows,
              "Tabela 1 - Síntese dos trabalhos relacionados e lacunas",
              "Fonte: elaborada pelos autores (2026).")

    add_text(
        doc,
        "Da análise consolidada emergem quatro lacunas que "
        "justificam a abordagem proposta: (i) nenhum trabalho "
        "encontrado aplica LLM com RAG diretamente ao COPSOQ II-Br "
        "nem ao arcabouço normativo brasileiro de SST; (ii) os "
        "pipelines mais avançados de anonimização em português "
        "brasileiro operam de forma desacoplada da geração de "
        "relatórios técnicos; (iii) os trabalhos sobre clustering "
        "temático de respostas abertas operam em pesquisas de "
        "opinião genéricas, não em respostas vinculadas a "
        "instrumento psicométrico validado de SST; e (iv) a "
        "literatura de geração automática de relatórios técnicos "
        "com LLM concentra-se em radiologia, prontuários e billing "
        "clínico, sem cobrir geração de Inventário de Riscos "
        "Psicossociais ou Plano de Ação 5W2H aderentes a uma norma "
        "regulatória de SST."
    )


def build_metodologia(doc):
    add_h1(doc, "5 METODOLOGIA")
    add_text(
        doc,
        "Este estudo adota abordagem aplicada, exploratória e de "
        "natureza tecnológica (GIL, 2002), com foco na concepção, "
        "na implementação e na avaliação quantitativa de cinco "
        "agentes baseados em LLM para a interpretação de respostas "
        "do COPSOQ II-Br e a produção dos documentos exigidos pela "
        "NR-01. A investigação se desdobrou em seis etapas "
        "encadeadas: revisão bibliográfica e análise documental do "
        "arcabouço normativo; modelagem do domínio e definição das "
        "fronteiras entre o motor analítico determinístico e os "
        "agentes generativos; implementação do sistema experimental "
        "Emotion Care; construção de gold standards específicos para "
        "cada agente; execução do protocolo de avaliação; e análise "
        "dos resultados em comparação com os trabalhos relacionados."
    )

    add_h2(doc, "5.1 Ambiente de desenvolvimento")
    add_text(
        doc,
        "A implementação e as experimentações foram conduzidas em "
        "estação de trabalho com processador AMD Ryzen 7 5700X "
        "(oito núcleos fisicos e dezesseis threads lógicos, 3,4 GHz "
        "de frequência base), 16 GB de memória RAM DDR4, GPU AMD "
        "Radeon RX 9060 XT (utilizada apenas para apresentação "
        "gráfica, sem treinamento local de modelos), armazenamento "
        "SSD NVMe e placa-mãe Gigabyte B550M DS3H AC. O sistema "
        "operacional adotado foi o Microsoft Windows 11 Pro versão "
        "10.0.26200 (arquitetura 64 bits), com PowerShell 5.1 para "
        "automação e Docker Desktop para empacotamento dos serviços "
        "auxiliares. O acesso aos LLM se dá via API REST autenticada "
        "da OpenAI (modelo gpt-4o-mini para geração e "
        "text-embedding-3-small para vetorização), com chave "
        "armazenada em arquivo .env e não versionada no repositório "
        "público do projeto, disponível em "
        "https://github.com/Alencar-png/emotion-care-tcc."
    )

    add_h2(doc, "5.2 Tecnologias adotadas")
    header = ["Camada", "Tecnologia", "Versão", "Propósito"]
    rows = [
        ["Linguagem (núcleo)", "Python", "3.11.5",
         "Implementação dos agentes e do motor analítico"],
        ["Linguagem (interface)", "TypeScript", "5",
         "Interface experimental"],
        ["Framework HTTP", "FastAPI", "0.111",
         "Endpoints dos agentes"],
        ["ORM", "SQLAlchemy", "2.0",
         "Mapeamento objeto-relacional"],
        ["Banco relacional", "PostgreSQL", "16",
         "Persistência de respostas"],
        ["Banco vetorial", "pgvector", "0.7", "Embeddings RAG"],
        ["Migrações", "Alembic", "1.13", "Versionamento de schema"],
        ["Orquestração LLM", "LangChain", "0.3",
         "Encadeamento prompt, modelo, validador"],
        ["LLM", "OpenAI gpt-4o-mini", "2024-07-18",
         "Geração dos agentes"],
        ["Embeddings", "text-embedding-3-small", "1536 dim.",
         "Vetorização do corpus normativo"],
        ["Containers", "Docker Compose", "v2",
         "Provisionamento local"],
        ["Avaliação", "scikit-learn, matplotlib", "ultimas",
         "Métricas e figuras"],
    ]
    add_table(doc, header, rows,
              "Tabela 2 - Tecnologias adotadas e respectivas versões",
              "Fonte: elaborada pelos autores (2026).")

    add_h2(doc, "5.3 Configurações do projeto")
    add_text(
        doc,
        "As configurações que governam o comportamento dos agentes "
        "foram consolidadas em três artefatos versionados: (a) "
        "limiares de anonimato com piso de quatro respondentes para "
        "qualquer score agregado e três por setor para recortes "
        "departamentais (SWEENEY, 2002); (b) configurações por caso "
        "de uso de LLM com temperaturas, limites de tokens de saída "
        "e modelo independentes por agente; (c) configuração do "
        "retriever RAG com chunk_size 1.600 caracteres, overlap 200, "
        "limiar de similaridade de cosseno 0,75 e top-k igual a 5."
    )

    add_h2(doc, "5.4 Fluxo da solução")
    add_text(
        doc,
        "O fluxo conceitual da solução proposta articula seis "
        "estágios entre a coleta da resposta do colaborador e a "
        "entrega dos documentos exigidos pela NR-01. As Figuras 1 a "
        "4 ilustram, respectivamente, a arquitetura do sistema, o "
        "user flow, o pipeline LLM/RAG e o processamento de uma "
        "resposta no agente narrador."
    )
    add_figure(doc, FIG / "fig_system_design.png",
               "Figura 1 - System Design: arquitetura conceitual do Emotion Care",
               "Fonte: elaborada pelos autores (2026).")
    add_figure(doc, FIG / "fig_user_flow.png",
               "Figura 2 - User Flow: colaborador respondente e gestor de SST",
               "Fonte: elaborada pelos autores (2026).")
    add_figure(doc, FIG / "fig_llm_pipeline.png",
               "Figura 3 - Pipeline LLM/RAG dos 5 agentes com validador anti-PII",
               "Fonte: elaborada pelos autores (2026).")
    add_figure(doc, FIG / "fig_llm_processing.png",
               "Figura 4 - Processamento de uma resposta no agente narrador do PGR",
               "Fonte: elaborada pelos autores (2026).")

    add_text(
        doc,
        "Estágio 1, Coleta: o colaborador acessa o COPSOQ II-Br por "
        "link personalizado e responde em escala Likert de cinco "
        "pontos. Estágio 2, Motor analítico determinístico: o "
        "agregador normaliza para a escala de zero a cem (com "
        "inversão de itens reverso) e calcula médias por dimensão e "
        "setor com classificação em três cores. Estágio 3, Filtro de "
        "anonimato (k-anonymity): bloqueia recortes abaixo do piso. "
        "Estágio 4, Agentes LLM: narrador, gerador 5W2H, copiloto "
        "NR-01 e analisador qualitativo. Estágio 5, Validador "
        "anti-PII: obrigatório em toda saída textual. Estágio 6, "
        "Devolução e exportação em PDF assinado."
    )

    add_h2(doc, "5.5 Construção dos gold standards")
    header = ["Agente", "Tamanho", "Natureza dos rótulos", "Origem"]
    rows = [
        ["Validador anti-PII", "1.025 amostras (445+ / 580-)",
         "8 categorias de PII com anotação span-level (offsets)",
         "Gerador determinístico (pii_gold_1000.py) com 100+ "
         "adversariais"],
        ["Narrador do PGR", "100 payloads sintéticos",
         "Inventário de referência determinístico para "
         "BERTScore/ROUGE",
         "Geração algorítmica (5 perfis x 4 portes x 5 seeds)"],
        ["Gerador 5W2H", "80 cenários cegos (20 por nível NR-01)",
         "Aderência à hierarquia NR-01 (4 níveis) por rubrica",
         "action_plan_gold_v3.py com descrição cega quanto à "
         "natureza do controle"],
        ["Copiloto NR-01 (RAG)", "40 perguntas-padrão",
         "Fonte verdadeira conhecida + RAGAS (faithfulness, answer "
         "relevancy, context relevancy)",
         "Curadoria sobre NR-01, NR-17, COPSOQ-II, ISO 45003, LGPD"],
        ["Analisador qualitativo", "150 respostas abertas",
         "7 temas pré-definidos + métricas de coerência (c_v)",
         "Geração algorítmica baseada em cenários reais; "
         "comparado contra BERTopic, k-means k=7 e LDA"],
    ]
    add_table(doc, header, rows,
              "Tabela 3 - Gold standards construídos para a avaliação dos agentes",
              "Fonte: elaborada pelos autores (2026).")

    add_h2(doc, "5.6 Métricas adotadas")
    add_text(
        doc,
        "Cada agente foi avaliado com métricas primárias justificadas "
        "pela função de utilidade do problema específico, "
        "complementadas por métricas secundárias diagnósticas, "
        "seguindo a recomendação de Sokolova e Lapalme (2009) de que a "
        "escolha da métrica deve ser orientada pelo perfil de erros "
        "tolerável no domínio de aplicação. As escolhas seguem três "
        "princípios metodológicos. Primeiro, para "
        "classificação com custos assimétricos entre falso positivo e "
        "falso negativo (validador anti-PII), Fβ=2 é preferível a F1; "
        "sob a Lei Geral de Proteção de Dados, vazar um dado pessoal "
        "(falso negativo) configura incidente reportável, ao passo que "
        "sobre-marcar um trecho (falso positivo) é inócuo, "
        "justificando peso quatro vezes maior ao recall (Manning, "
        "Raghavan e Schütze, 2008, cap. 8). Segundo, para geração de "
        "texto técnico, métricas de similaridade contra referência "
        "(BERTScore e ROUGE) são complementadas por métricas de "
        "domínio: faithfulness ao payload, taxa de alucinação "
        "regulatória contra base curada de itens da NR-01, NR-17, "
        "COPSOQ-II e ISO 45003, e avaliação humana cega Likert "
        "(Voinea et al., 2024). Terceiro, para Geração Aumentada por "
        "Recuperação em domínio regulado, o framework RAGAS (Es et "
        "al., 2024) é o padrão atual, cobrindo simultaneamente a "
        "etapa de recuperação (precision@k, recall@k, MRR) e a etapa "
        "de geração (faithfulness, answer relevancy, context "
        "relevancy)."
    )
    add_text(
        doc,
        "Em consonância com esses princípios, o validador anti-PII é "
        "avaliado em dois níveis: documento (presença binária de PII) "
        "e span (entidade por entidade, com offsets de caractere). "
        "Em ambos, a métrica primária é Fβ=2 e a métrica secundária é "
        "F1, com matriz de confusão por categoria e taxa de vazamento "
        "categórica (fração de documentos que contêm PII de uma "
        "categoria e tiveram pelo menos uma entidade dessa categoria "
        "escapando da detecção). A curva ROC e o AUC foram "
        "deliberadamente excluídos: o validador é classificador "
        "determinístico (expressões regulares mais heurística), sem "
        "score contínuo nem threshold variável; reportar AUC seria "
        "tecnicamente incorreto."
    )
    add_text(
        doc,
        "O narrador do PGR é avaliado por (i) BERTScore F1 contra "
        "inventário de referência gerado deterministicamente a partir "
        "do payload, (ii) ROUGE-1 e ROUGE-L F1 contra a mesma "
        "referência, (iii) faithfulness ao payload (fração de "
        "sentenças com afirmação numérica ancorada nos campos "
        "originais), (iv) taxa de alucinação regulatória e (v) "
        "conformidade estrutural (sete seções, faixa 600 a 900 "
        "palavras, ausência de PII). O gerador 5W2H reporta acurácia "
        "sobre a hierarquia NR-01 com matriz de confusão 4x4 e "
        "intervalo de confiança 95%, taxa de aderência andragógica "
        "(presença explícita de princípios de Knowles), diversidade "
        "de ações (entropia de Shannon e type-token ratio) e taxa "
        "de alucinação regulatória. O copiloto NR-01 reporta "
        "precision@k, recall@k e MRR para k em {1, 3, 5}, além das "
        "três métricas RAGAS e taxa de alucinação regulatória. "
        "Rubricas Likert para specificity, implementability e "
        "clareza percebida foram pré-registradas (módulo "
        "likert_rubrics.py) e amostras estratificadas exportadas "
        "para aplicação por dois revisores independentes com kappa "
        "de Cohen ponderado; a aplicação humana fica como trabalho "
        "futuro (Capítulo 9). O "
        "analisador qualitativo reporta ARI, NMI, homogeneity, "
        "completeness, V-measure, F1 macro e weighted, coerência "
        "tópica c_v (gensim) com c_npmi de complemento, e stability "
        "across seeds (cinco execuções com sementes distintas). "
        "Três baselines são comparados: BERTopic, k-means k=7 sobre "
        "embeddings text-embedding-3-small, e LDA com sete tópicos."
    )
    add_text(
        doc,
        "Todas as métricas pontuais reportadas no Capítulo 6 são "
        "acompanhadas de intervalos de confiança 95% obtidos por "
        "bootstrap não-paramétrico (Efron e Tibshirani, 1993) com "
        "mil reamostragens com reposição. Comparações entre o "
        "presente trabalho e a literatura seguem teste pareado "
        "quando os datasets coincidem (McNemar com correção de "
        "continuidade) e teste bootstrap entre datasets distintos "
        "(reportando apenas diferença descritiva e intervalo de "
        "confiança individual de cada lado). Métricas operacionais "
        "(latência p50, p95 e p99, tokens de entrada e saída, "
        "custo estimado em dólar por inferência e taxa de erro ou "
        "timeout em chamadas a API) são capturadas em logs "
        "estruturados para todos os agentes."
    )
    add_table(doc,
              ["Agente", "Métrica primária 1",
               "Métrica primária 2", "Métrica primária 3"],
              [["Validador anti-PII",
                "Fβ=2 doc-level",
                "Fβ=2 macro span-level",
                "Taxa de vazamento por categoria"],
               ["Narrador do PGR",
                "BERTScore F1 vs. referência",
                "Faithfulness ao payload",
                "Taxa de alucinação regulatória"],
               ["Gerador 5W2H",
                "Acurácia 4x4 (gold cego)",
                "Aderência andragógica (Knowles)",
                "Taxa de alucinação regulatória"],
               ["Copiloto NR-01 (RAG)",
                "Faithfulness (RAGAS)",
                "Answer Relevancy (RAGAS)",
                "Taxa de alucinação regulatória"],
               ["Analisador qualitativo",
                "Coerência tópica c_v",
                "ARI vs. baselines (BERTopic/k-means/LDA)",
                "Stability across seeds (σ ARI)"]],
              "Quadro 1 - Aparato metodológico: métricas primárias por agente",
              "Fonte: elaborado pelos autores (2026). Métricas secundárias e operacionais (latência, custo, "
              "tokens, conformidade estrutural) são reportadas em conjunto no Capítulo 6.")

    add_h2(doc, "5.7 Garantias de segurança e privacidade do experimento")
    add_text(
        doc,
        "Todos os dados utilizados na avaliação são sintéticos, "
        "gerados algoritmicamente para representar cenários "
        "ocupacionais plausíveis sem corresponder a indivíduos "
        "reais. Não houve coleta de dados de pessoas físicas, "
        "dispensando submissão a comitê de ética para esta etapa. "
        "A validação com dados reais e prevista como trabalho "
        "futuro, condicionada a aprovação do Comitê de Ética em "
        "Pesquisa."
    )


def build_resultados(doc):
    add_h1(doc, "6 RESULTADOS")
    add_text(
        doc,
        "Este capítulo apresenta os resultados quantitativos "
        "obtidos na avaliação dos cinco agentes baseados em LLM. A "
        "discussão comparativa com os trabalhos relacionados e "
        "apresentada no Capítulo 7."
    )

    add_h2(doc, "6.1 Validador anti-PII")
    add_text(
        doc,
        "A avaliação do validador anti-PII utiliza um gold "
        "standard ampliado com 1.025 casos "
        "rotulados (445 positivos e 580 negativos) com anotação de "
        "spans em offsets de caractere, cobrindo oito categorias de "
        "PII: e-mail, CPF, CNPJ, telefone, matrícula, cargo "
        "identificador, título com prenome e nome composto. Os 485 "
        "spans positivos somam mais que os 445 documentos positivos "
        "porque vinte casos foram desenhados como multi-categoria "
        "(por exemplo, nome próprio + CPF + e-mail em um mesmo "
        "documento), cada um contendo três spans rotulados, "
        "totalizando 425 documentos com um span único e 20 "
        "documentos com três spans (425 + 60 = 485 spans em 445 "
        "documentos). Os spans distribuem-se com aproximadamente 50 "
        "a 80 amostras por categoria. Foram incluídos 100 casos "
        "adversariais (capitalização institucional como Plano de "
        "Ação, Recursos Humanos, Universidade Federal de Alagoas; "
        "sequências numéricas como ISBN e número de série) "
        "projetados para induzir falsos positivos em heurísticas "
        "ingênuas. As métricas primárias passaram a ser Fβ=2 e "
        "Fβ=2 macro span-level; AUC e curva ROC foram removidas por "
        "serem tecnicamente inválidas para classificador "
        "determinístico sem score contínuo, conforme justificado na "
        "Seção 5.6."
    )
    add_text(
        doc,
        "A Tabela 4 apresenta os resultados no nível de documento, "
        "com intervalos de confiança 95% obtidos por bootstrap "
        "não-paramétrico (n=1.000 reamostragens)."
    )
    add_table(doc,
              ["Métrica", "Ponto", "IC 95%"],
              [["Verdadeiros positivos (TP)", "429", "/"],
               ["Verdadeiros negativos (TN)", "580", "/"],
               ["Falsos positivos (FP)", "0", "/"],
               ["Falsos negativos (FN)", "16", "/"],
               ["Precisão", "1,0000", "[1,0000; 1,0000]"],
               ["Recall", "0,9640", "[0,9459; 0,9804]"],
               ["F1-score", "0,9817", "[0,9722; 0,9901]"],
               ["Fβ=2 (primária)", "0,9710", "[0,9563; 0,9843]"],
               ["Latência p50", "0,020 ms", "/"],
               ["Latência p95", "0,040 ms", "/"],
               ["Latência p99", "0,062 ms", "/"]],
              "Tabela 4 - Resultados do validador anti-PII no nível de documento (n = 1.025)",
              "Fonte: dados da pesquisa, execução em 28/05/2026.")

    add_figure(doc, FIG / "pii_v3_confusion_doc.png",
               "Figura 5 - Matriz de confusão doc-level do validador anti-PII (n=1.025)",
               "Fonte: dados da pesquisa (2026).")

    add_text(
        doc,
        "A precisão de 1,0000 indica zero falso positivo no "
        "nível de documento sobre 580 textos negativos, mesmo "
        "incluindo casos adversariais como Plano de Ação 5W2H, "
        "Norma Regulamentadora, Banco Central do Brasil e ISBN "
        "978-85-1234-567-8. Esse resultado decorre da lista de "
        "tokens institucionais e do mascaramento ordenado CPF, "
        "CNPJ, telefone implementados no validador determinístico "
        "(Seção 5.3.1). O recall de 0,9640 mostra que 16 dos 445 "
        "documentos positivos tiveram pelo menos uma categoria de "
        "PII não detectada, reflexo das fragilidades residuais nas "
        "categorias telefone (8 casos com formatos internacionais "
        "menos comuns) e matrícula (4 casos com gatilhos textuais "
        "raros). A alternativa intl_split do regex de telefone "
        "cobre o padrão +55DDD 9 NNNN-NNNN e os gatilhos "
        "pós-numéricos do regex de matrícula (sistema de ponto, "
        "folha de pagamento, ADP, Senior, SAP, AD) garantem "
        "recall de 0,9200 nessa categoria. A métrica primária "
        "Fβ=2 = 0,9710 (IC95% [0,9563; 0,9843]) reflete "
        "corretamente o trade-off precisão-recall, penalizando o "
        "vazamento conforme a função de utilidade do problema sob "
        "LGPD."
    )

    add_text(
        doc,
        "A Tabela 5 apresenta as métricas span-level (entidade por "
        "entidade) por categoria, com match parcial (sobreposição "
        "de offsets maior ou igual a 50% do menor span entre "
        "predito e gold). Essa é a granularidade utilizada por "
        "trabalhos comparáveis em NER e de-identification."
    )
    add_table(doc,
              ["Categoria", "TP", "FP", "FN", "Suporte",
               "Precisão", "Recall", "F1", "Fβ=2"],
              [["e-mail", "80", "0", "0", "80",
                "1,0000", "1,0000", "1,0000", "1,0000"],
               ["CPF", "75", "0", "0", "75",
                "1,0000", "1,0000", "1,0000", "1,0000"],
               ["CNPJ", "50", "0", "0", "50",
                "1,0000", "1,0000", "1,0000", "1,0000"],
               ["Telefone", "47", "0", "8", "55",
                "1,0000", "0,8545", "0,9216", "0,8801"],
               ["Matrícula", "46", "0", "4", "50",
                "1,0000", "0,9200", "0,9583", "0,9350"],
               ["Cargo identificador", "50", "0", "0", "50",
                "1,0000", "1,0000", "1,0000", "1,0000"],
               ["Título + nome", "50", "0", "0", "50",
                "1,0000", "1,0000", "1,0000", "1,0000"],
               ["Nome composto", "69", "0", "6", "75",
                "1,0000", "0,9200", "0,9583", "0,9350"],
               ["Macro span-level", "/", "/", "/", "/",
                "1,0000", "0,9474", "0,9700", "0,9688"],
               ["Macro Fβ=2 (IC95%)", "/", "/", "/", "/",
                "/", "/", "/", "[0,955; 0,982]"]],
              "Tabela 5 - Métricas span-level por categoria (match parcial >=50%, n=485 spans)",
              "Fonte: dados da pesquisa (2026).")

    add_figure(doc, FIG / "pii_v3_span_metrics.png",
               "Figura 6 - Métricas span-level por categoria",
               "Fonte: dados da pesquisa (2026).")

    add_text(
        doc,
        "A Tabela 5b apresenta a taxa de vazamento por categoria, "
        "métrica de domínio específica: fração de documentos que "
        "contêm PII de uma categoria e tiveram pelo menos uma "
        "entidade dessa categoria escapando da detecção. É o "
        "número operacional que um Encarregado de Dados (DPO) "
        "examina para julgar o risco residual da arquitetura."
    )
    add_table(doc,
              ["Categoria", "Docs com PII", "Docs com vazamento",
               "Taxa de vazamento"],
              [["e-mail", "80", "0", "0,0000"],
               ["CPF", "75", "0", "0,0000"],
               ["CNPJ", "50", "0", "0,0000"],
               ["Telefone", "55", "8", "0,1455"],
               ["Matrícula", "50", "4", "0,0800"],
               ["Cargo identificador", "50", "0", "0,0000"],
               ["Título + nome", "50", "0", "0,0000"],
               ["Nome composto", "75", "6", "0,0800"]],
              "Tabela 5b - Taxa de vazamento por categoria (validador anti-PII)",
              "Fonte: dados da pesquisa (2026).")

    add_text(
        doc,
        "Pontos de fragilidade residuais identificados nesta "
        "avaliação. (i) Categoria telefone com 8 falsos negativos "
        "span-level (recall 0,8545): correspondem a telefones "
        "internacionais com formatação ambígua e a padrões de DDI "
        "raros em PT-BR ocupacional. (ii) Categoria matrícula "
        "com 4 falsos negativos (recall 0,9200): correspondem a "
        "templates em que o número de matrícula aparece sem "
        "qualquer gatilho textual prévio ou pós-fixado dentro do "
        "conjunto reconhecido. (iii) Categoria nome composto "
        "com 6 falsos negativos (recall 0,9200): predominantemente "
        "nomes próprios escritos sem acentos. As categorias "
        "e-mail, CPF, CNPJ, cargo identificador e título + nome "
        "atingiram F1 = 1,0000, sem qualquer escape detectável "
        "no gold. As mitigações propostas para os casos "
        "restantes (camada secundária de NER fina via BERTimbau "
        "ou Llama-3 acionada somente em caso de ausência de "
        "match regex) estão documentadas no Capítulo 9 como "
        "trabalho futuro."
    )
    add_text(
        doc,
        "Comparação descritiva com Schiezaro et al. (2026). O "
        "presente trabalho atinge F1 doc-level de 0,9699 (IC95% "
        "[0,9577; 0,9811]) sobre 1.025 amostras sintéticas, enquanto "
        "Schiezaro et al. (2026) reportam F1 de 0,9270 sobre 2.962 "
        "registros clínicos reais (BERTimbau-leNER). A diferença de "
        "ponto é +0,0429 a favor do presente trabalho, porém a "
        "comparação é descritiva, não inferencial: datasets, "
        "instrumentação e domínio diferem (sintético ocupacional vs. "
        "clínico real). Bootstrap pareado entre datasets distintos "
        "indica ordem de grandeza compatível, sem sustentar a "
        "afirmação de superioridade metodológica. A vantagem real "
        "da arquitetura presente, em relação a abordagens baseadas "
        "em BERT/LLM, está no perfil de latência (p99 inferior a "
        "0,1 ms vs. dezenas a centenas de milissegundos de "
        "inferência neural), o que viabiliza o validador como "
        "guard-rail síncrono em pipelines de geração."
    )

    add_h2(doc, "6.2 Narrador do PGR")
    add_text(
        doc,
        "A avaliação do narrador do PGR utiliza o aparato "
        "metodológico justificado na Seção 5.6: gold de 100 "
        "payloads (cinco perfis epidemiológicos "
        "saudável/alerta_um/alerta_multi/critico_focal/critico_multi "
        "vezes quatro portes vezes cinco seeds), inventário de "
        "referência gerado deterministicamente a partir do payload "
        "(módulo reference_inventory_generator.py) e métricas "
        "primárias por BERTScore F1, ROUGE-1, ROUGE-L, faithfulness "
        "ao payload e taxa de alucinação regulatória. O re-prompting "
        "iterativo está habilitado, conforme a configuração de "
        "produção descrita na Seção 5.3.2. Todas as métricas pontuais "
        "reportam IC95% por bootstrap (n=1.000)."
    )
    add_table(doc,
              ["Métrica", "Valor", "IC 95%"],
              [["BERTScore F1 vs. referência (primária)", "0,7623",
                "[0,7590; 0,7660]"],
               ["ROUGE-1 F1", "0,4382", "/"],
               ["ROUGE-L F1", "0,2410", "/"],
               ["Faithfulness ao payload (primária)", "0,9316",
                "[0,9153; 0,9466]"],
               ["Taxa de alucinação regulatória (primária)", "0,0000",
                "0 alucinações observadas"],
               ["Presença das 7 seções", "1,0000",
                "100/100"],
               ["Ausência de PII na saída", "1,0000",
                "100/100"],
               ["Contagem na faixa 600 a 900 palavras", "0,6300",
                "63/100 (com re-prompting iterativo)"],
               ["Latência p50", "47,28 s", "/"],
               ["Latência p95", "72,51 s", "/"],
               ["Latência p99", "82,29 s", "/"]],
              "Tabela 6 - Resultados do narrador do PGR (gold v3, n=100)",
              "Fonte: dados da pesquisa, execução em 28/05/2026. BERTScore "
              "computado com bert-base-multilingual-cased (fallback do "
              "neuralmind/bert-base-portuguese-cased, layer=9).")
    add_text(
        doc,
        "BERTScore F1 de 0,7623 (IC95% [0,7590; 0,7660]) indica "
        "alta similaridade semântica entre o texto gerado e o "
        "inventário de referência determinístico, valor "
        "compatível com a faixa observada em geração condicional "
        "no domínio técnico (VOINEA et al., 2024, reportam "
        "BERTScore 0,8054 para laudos radiológicos com Llama-3 "
        "fine-tuned, em outro domínio). ROUGE-1 F1 de 0,4382 "
        "confirma sobreposição lexical moderada de unigramas, "
        "enquanto o ROUGE-L F1 de 0,2410 expõe que as sequências "
        "mais longas (frases inteiras) divergem do gerador "
        "determinístico, o que é esperado: o LLM parafraseia o "
        "conteúdo factual em vez de copiar a estrutura sintática "
        "do template de referência."
    )
    add_text(
        doc,
        "Faithfulness ao payload de 0,9316 (IC95% [0,9153; 0,9466]) "
        "indica que aproximadamente 93% das sentenças com "
        "afirmação numérica na saída têm pelo menos um número "
        "ancorado nos campos originais do payload (escores, "
        "porcentagens, contagens, dimensões classificadas). Esse "
        "patamar reflete o efeito do prompt revisado com regra "
        "explícita de ancoragem numérica e distribuição obrigatória "
        "por seção. Taxa de alucinação regulatória de 0,0000 "
        "confirma que o narrador, no contexto deste experimento, "
        "cita apenas itens normativos verificáveis contra a base "
        "curada, sem produzir referências inexistentes."
    )
    add_text(
        doc,
        "A conformidade na faixa de palavras (600 a 900) é de "
        "0,6300 com o prompt do narrador combinado a re-prompting "
        "iterativo (até duas iterações) e regra explícita de "
        "distribuição mínima de palavras por seção. Os 37 textos "
        "fora da faixa correspondem majoritariamente a payloads de "
        "pequeno porte com poucos riscos elevados, em que a "
        "narrativa naturalmente comprime a contagem; isso motiva "
        "trabalho futuro com instrução condicional ao porte da "
        "organização e à criticidade do payload."
    )
    add_text(
        doc,
        "A taxa de presença das sete seções obrigatórias é 1,0000, "
        "indicando que mesmo cenários saudáveis sem dimensão "
        "Amarela ou Vermelha geram a seção \"riscos intermediários\" "
        "com a indicação explícita de ausência de risco no nível "
        "correspondente. A taxa de ausência de PII na saída atinge "
        "igualmente 1,0000 nos 100 textos gerados, evidenciando "
        "que o narrador não reintroduz dados pessoais no "
        "inventário a partir do payload anônimo. A latência "
        "mediana de 47,28 segundos (p95 = 72,51 s; p99 = 82,29 s) "
        "reflete o custo do re-prompting iterativo, mas permanece "
        "compatível com uso assíncrono na geração do Inventário."
    )

    add_h2(doc, "6.3 Gerador de Plano de Ação 5W2H")
    add_text(
        doc,
        "A avaliação do gerador 5W2H é conduzida sobre o gold "
        "action_plan_gold_v3.py, com 80 cenários distribuídos "
        "uniformemente nos quatro níveis da hierarquia NR-01 "
        "(vinte por nível). A descrição de cada cenário é cega "
        "quanto à natureza do controle apropriado, eliminando "
        "vazamento de rótulo, e a rubrica está pronta para "
        "aplicação por dois revisores independentes."
    )
    add_table(doc,
              ["Métrica", "Valor", "IC 95%"],
              [["Acurácia 4 níveis (primária)", "0,7875", "[0,7000; 0,8750]"],
               ["Macro F1 (4 níveis)", "0,7828", "/"],
               ["Macro precisão", "0,7993", "/"],
               ["Macro recall", "0,7875", "/"],
               ["Taxa de aderência andragógica (Knowles)", "0,9250",
                "74/80 planos"],
               ["Taxa de alucinação regulatória", "0,0000",
                "76 citações válidas / 76"],
               ["Type-Token Ratio (diversidade de ações)", "0,0927", "/"],
               ["Entropia Shannon normalizada", "1,0000", "/"],
               ["Latência p50", "7,1 s", "/"],
               ["Latência p95", "12,61 s", "/"],
               ["Latência p99", "15,08 s", "/"]],
              "Tabela 7 - Resultados do gerador 5W2H (gold v3, n=80, descrição cega)",
              "Fonte: dados da pesquisa, execução em 28/05/2026.")
    add_figure(doc, FIG / "action_plan_confusion.png",
               "Figura 7 - Matriz de confusão 4x4 da hierarquia de controle NR-01",
               "Fonte: dados da pesquisa (2026).")
    add_text(
        doc,
        "A Tabela 7b detalha precisão, recall e F1 por nível da "
        "hierarquia, revelando padrão de viés do agente:"
    )
    add_table(doc,
              ["Nível NR-01", "Precisão", "Recall", "F1", "Suporte",
               "Observação"],
              [["1-Eliminação", "0,9290", "0,6500", "0,7650", "20",
                "Precisão alta, recall moderado"],
               ["2-Substituição", "0,6670", "0,6000", "0,6320", "20",
                "Recall recuperado com exemplos few-shot "
                "contrastando Nível 1 versus Nível 2"],
               ["3-Controle Organizacional", "0,6920", "0,9000",
                "0,7830", "20", "Recall alto, precisão moderada"],
               ["4-Controle Individual", "0,9090", "1,0000",
                "0,9520", "20", "Acerto quase perfeito"]],
              "Tabela 7b - Métricas por nível NR-01 (gold v3, n=80)",
              "Fonte: dados da pesquisa (2026).")
    add_text(
        doc,
        "A matriz de confusão (Figura 7) é equilibrada entre os "
        "quatro níveis. O agente acerta vinte dos vinte cenários "
        "de Nível 4 (Controle Individual), treze dos vinte de "
        "Nível 1 (Eliminação), doze dos vinte de Nível 2 "
        "(Substituição) e dezoito dos vinte de Nível 3 (Controle "
        "Organizacional). A árvore de decisão de quatro perguntas, "
        "as definições essenciais por nível e os seis exemplos "
        "few-shot do prompt (Seção 5.3.3), incluindo três casos "
        "contrastando explicitamente Nível 1 versus Nível 2 e "
        "Nível 2 versus Nível 3, sustentam recall de 0,6000 no "
        "Nível 2, valor consideravelmente mais alto do que o "
        "patamar de 0,3500 característico de prompts sem "
        "contraste explícito entre os níveis adjacentes."
    )
    add_text(
        doc,
        "Outras métricas do agente preservam comportamento "
        "satisfatório. A aderência andragógica é de 0,9250, com 74 "
        "dos 80 planos referenciando explicitamente ao menos um "
        "dos cinco princípios de Knowles (autonomia, experiência "
        "prévia, prontidão, orientação para aplicação, motivação "
        "interna), satisfazendo o requisito do prompt. A taxa de "
        "alucinação regulatória é de 0,0000: as setenta e seis "
        "citações normativas extraídas dos planos foram todas "
        "validadas contra a base curada NR-01/NR-17/COPSOQ/ISO 45003. "
        "A diversidade de ações (entropia normalizada 1,0000) "
        "confirma que cada plano é único, sem mode collapse "
        "lexical. A latência mediana de 7,1 segundos cresce para "
        "12,6 segundos no p95 quando o cenário envolve múltiplos "
        "riscos, dentro do envelope aceitável para uso assíncrono. "
        "A próxima etapa de avaliação inclui aplicação "
        "Likert humana das rubricas specificity e implementability "
        "pelos dois revisores; as amostras estratificadas estão "
        "exportadas em metrics/likert_samples/."
    )

    add_h2(doc, "6.4 Copiloto NR-01 com RAG")
    add_text(
        doc,
        "A avaliação do copiloto utiliza o framework RAGAS "
        "(Es et al., 2024), que cobre tanto a "
        "etapa de recuperação (precision@k, recall@k, MRR) quanto "
        "a etapa de geração (faithfulness, answer relevancy e "
        "context relevancy). Acrescenta-se ainda a taxa de "
        "alucinação regulatória contra base curada de itens da "
        "NR-01, NR-17, COPSOQ-II e ISO 45003 (módulo "
        "regulatory_lookup.py). O gold contém 40 perguntas-padrão "
        "cobrindo essas normas, LGPD e a política interna de "
        "k-anonymity. O limiar de similaridade cosseno mantido em "
        "0,40 (não no valor de produção 0,75): a base de "
        "conhecimento contém apenas excertos representativos por "
        "documento, e thresholds mais altos rejeitavam recuperações "
        "válidas em corpus reduzido. A expansão do corpus para "
        "documentos normativos completos é trabalho futuro e a "
        "Seção 9 detalha o plano de ingestão. Todas as métricas "
        "pontuais reportam IC95% por bootstrap (n=1.000)."
    )
    add_table(doc,
              ["Métrica", "Valor", "IC 95%"],
              [["Precision@1 (recuperação)", "0,6500",
                "[0,5000; 0,8000]"],
               ["Precision@3", "0,2750", "/"],
               ["Precision@5", "0,1650", "/"],
               ["Recall@1", "0,6500", "/"],
               ["Recall@3", "0,8250", "/"],
               ["Recall@5", "0,8250", "[0,7000; 0,9250]"],
               ["MRR", "0,7271", "[0,6125; 0,8500]"],
               ["Faithfulness (primária RAGAS)", "0,1514",
                "[0,0915; 0,2219]"],
               ["Answer relevancy (primária RAGAS)", "0,6116",
                "[0,5547; 0,6706]"],
               ["Context relevancy", "0,5508", "[0,4242; 0,6925]"],
               ["Taxa de alucinação regulatória", "0,0000",
                "0/19 citações"],
               ["Latência p50", "6,12 s", "/"],
               ["Latência p95", "8,80 s", "/"],
               ["Latência p99", "10,42 s", "/"]],
              "Tabela 8 - Resultados do copiloto NR-01 com RAGAS (n=40, threshold 0,40, corpus integral)",
              "Fonte: dados da pesquisa, execução em 28/05/2026. Corpus indexado: 596 chunks de NR-01, NR-17, COPSOQ, ISO 45003, LGPD e política de k-anonymity.")
    add_text(
        doc,
        "A recuperação atinge Precision@1 = 0,6500 (IC95% [0,500; "
        "0,800]) e Recall@5 = 0,8250 (IC95% [0,700; 0,925]): em "
        "82,5% das 40 perguntas, o documento-fonte correto está "
        "entre os cinco primeiros resultados retornados; em 65% "
        "delas ele já é o primeiro. O MRR de 0,7271 indica que a "
        "fonte correta aparece em torno da posição ranqueada 1,4, "
        "em média. A indexação contempla os textos integrais de "
        "NR-01, NR-17 e LGPD (somando 574 chunks de 600 "
        "caracteres com overlap de 120) e excertos curados de "
        "COPSOQ-II, ISO 45003 e da política de k-anonymity, "
        "totalizando 596 chunks indexados (Seção 5.3.4)."
    )
    add_text(
        doc,
        "O achado central, e o ponto de fragilidade mais sensível "
        "do copiloto, é a faithfulness de 0,1514 (IC95% "
        "[0,0915; 0,2219]): apenas 15% das sentenças geradas pelo "
        "agente possuem suporte semântico direto nos chunks "
        "recuperados (cosseno entre embedding da sentença e do "
        "melhor chunk maior ou igual a 0,75). O valor está bem "
        "abaixo dos patamares praticados pela literatura RAG (Es "
        "et al., 2024, reportam faithfulness 0,75 a 0,85 para QA "
        "assistido em domínio biomédico). A answer relevancy de "
        "0,6116 mostra que as respostas, ainda que pouco "
        "ancoradas, endereçam corretamente o tópico da pergunta. "
        "A context relevancy de 0,5508 indica que aproximadamente "
        "55% dos chunks recuperados têm cosseno com a pergunta "
        "maior ou igual a 0,55, valor compatível com o tamanho do "
        "corpus."
    )
    add_text(
        doc,
        "A taxa de alucinação regulatória de 0,0000 sobre dezenove "
        "citações extraídas das respostas é o sinal operacional "
        "mais relevante do agente: nenhuma referência normativa "
        "produzida nesta avaliação aponta para item inexistente. "
        "O resultado decorre da combinação entre o prompt com "
        "regra de ouro de abstenção, a base de validação "
        "regulatória com lookup hierárquico (que aceita subitens "
        "como prefixos válidos de itens curados) e o corpus "
        "integral indexado, que reduz a tentação de produção "
        "paramétrica de itens. Faithfulness baixa ainda motiva "
        "trabalhos futuros descritos no Capítulo 9: reranking "
        "pós-retrieval via cross-encoder, ou substituição da "
        "etapa generativa por modelos open-source com melhor "
        "disciplina de citação, como descrito no benchmark da "
        "Seção 7.8."
    )

    add_h2(doc, "6.5 Analisador qualitativo")
    add_text(
        doc,
        "A avaliação do analisador qualitativo combina cinco "
        "elementos metodológicos: (i) coerência "
        "tópica c_v (via gensim) e c_npmi como métricas primárias, "
        "padrão em topic modeling; (ii) F1 macro e weighted em "
        "lugar de apenas macro, para refletir corretamente a "
        "frequência dos temas; (iii) stability across seeds com "
        "cinco execuções (média e desvio-padrão de ARI, NMI e c_v); "
        "(iv) comparação direta contra três baselines (k-means k=7 "
        "sobre embeddings text-embedding-3-small, LDA com sete "
        "tópicos e BERTopic com UMAP+HDBSCAN); e (v) IC95% por "
        "bootstrap. A avaliação foi conduzida sobre 150 respostas "
        "sintéticas distribuídas em 7 temas (Sobrecarga, "
        "Reconhecimento, Ambiente físico, Comunicação, Carreira, "
        "Liderança e Outros), com temperatura do LLM fixada em "
        "zero e prompt enrijecido com fronteiras semânticas "
        "explícitas para cada tema, conforme detalhado na Seção "
        "5.3.5."
    )
    add_table(doc,
              ["Método", "ARI", "NMI", "c_v", "macro F1",
               "weighted F1"],
              [["LLM (média 5 seeds, temp=0)",
                "0,1498 ± 0,0069",
                "0,3599 ± 0,0155",
                "0,4724 ± 0,1057",
                "0,3649 ± /",
                "0,3662 ± /"],
               ["k-means (k=7, embedding)",
                "0,3681", "0,5592", "0,3517",
                "0,5332", "0,5345"],
               ["LDA (k=7)",
                "0,1625", "0,3222", "0,3460",
                "0,2984", "0,3007"],
               ["BERTopic (UMAP+HDBSCAN)",
                "0,1133", "0,4139", "0,3552",
                "0,3999", "0,4012"]],
              "Tabela 9 - Analisador qualitativo, LLM versus baselines (n=150, 7 temas)",
              "Fonte: dados da pesquisa, execução em 28/05/2026.")
    add_text(
        doc,
        "A comparação contra baselines revela padrão duplo. Em "
        "concordância categorial (ARI), o LLM atinge 0,1498 "
        "(desvio-padrão 0,0069 sobre cinco seeds, temperatura "
        "zero), abaixo do k-means (0,3681) que opera sobre os "
        "mesmos embeddings text-embedding-3-small e do LDA "
        "(0,1625), mas acima do BERTopic (0,1133). Em coerência "
        "tópica c_v (métrica primária de topic modeling), o LLM "
        "lidera com 0,4724 (desvio-padrão 0,1057) contra "
        "0,3517 do k-means, 0,3552 do BERTopic e 0,3460 do LDA. "
        "Essa dissociação entre concordância categorial e "
        "coerência tópica é coerente com o achado de Röder, Both "
        "e Hinneburg (2015) de que c_v alto não implica alta "
        "concordância: o LLM produz clusters semanticamente "
        "coesos (palavras-chave fortemente correlacionadas no "
        "co-occurrence-graph) mas com fronteiras temáticas que "
        "diferem do gold rotulado, sobretudo em \"Carreira\" "
        "(frequentemente fundido com \"Reconhecimento\") e "
        "\"Outros\"."
    )
    add_text(
        doc,
        "À luz desse padrão duplo, a escolha do LLM como método "
        "principal mantém-se justificada pelas vantagens "
        "operacionais (nomes descritivos por cluster sem etapa "
        "adicional de rotulação, controle explícito do número "
        "máximo de clusters via prompt, descrição semântica do "
        "tema na mesma chamada) e pela liderança em coerência "
        "tópica c_v, ainda que com perda em ARI frente ao "
        "k-means. O Capítulo 9 propõe três caminhos para reduzir "
        "essa perda: (i) ensemble LLM+k-means com voto "
        "majoritário, combinando a coerência semântica do LLM "
        "com a fronteira categorial do k-means; (ii) gold "
        "ampliado para 500 respostas e definição operacional "
        "mais clara dos temas Carreira e Outros; (iii) prompt "
        "menos prescritivo que permita o LLM redefinir suas "
        "fronteiras temáticas dentro de um limite máximo de "
        "clusters."
    )
    add_text(
        doc,
        "A Tabela 9b detalha o desempenho por tema do LLM no último "
        "seed avaliado, evidenciando que os temas Carreira e "
        "Outros permanecem como pontos de fragilidade do mapeamento "
        "cluster → tema verdadeiro."
    )
    add_table(doc,
              ["Tema", "Precisão", "Recall", "F1", "Suporte"],
              [["Sobrecarga", "0,7143", "0,6818", "0,6977", "22"],
               ["Reconhecimento", "1,0000", "0,5909", "0,7429", "22"],
               ["Ambiente físico", "0,9091", "0,9524", "0,9302", "21"],
               ["Comunicação", "0,9333", "0,6364", "0,7568", "22"],
               ["Carreira", "0,0000", "0,0000", "0,0000", "22"],
               ["Liderança", "0,4500", "1,0000", "0,6207", "20"],
               ["Outros", "0,0000", "0,0000", "0,0000", "21"]],
              "Tabela 9b - Desempenho por tema do LLM (seed=2, último)",
              "Fonte: dados da pesquisa (2026).")
    add_text(
        doc,
        "Sobrecarga, Reconhecimento, Ambiente físico e Comunicação "
        "são identificados com F1 entre 0,70 e 0,93. Liderança "
        "apresenta alto recall e precisão "
        "moderada, indicando classificação excessiva quando o tema "
        "verdadeiro é Sobrecarga ou Reconhecimento correlato. "
        "Carreira e Outros permanecem com F1 próximo de zero "
        "porque os clusters gerados foram mapeados, via "
        "similaridade semântica de embeddings, a outros temas "
        "vizinhos, indicando fronteira semântica difusa para "
        "categorias residuais. ARI agregado de 0,1498 (cinco "
        "seeds, temperatura zero) expressa concordância "
        "modesta acima do acaso; somado ao colapso parcial dos "
        "temas Carreira e Outros, configura cenário de "
        "concordância baixa com colapso parcial, em consonância "
        "com a literatura sobre dificuldade de fronteira "
        "semântica em respostas abertas curtas (Jiang, Liu e "
        "Fisher, 2026). O LLM lidera, contudo, na coerência "
        "tópica c_v (0,4724 contra 0,3517 do k-means), métrica "
        "primária de topic modeling, evidenciando que os "
        "clusters gerados são semanticamente coesos mesmo "
        "quando a fronteira categorial diverge do gold."
    )

    add_h2(doc, "6.6 Métricas operacionais do sistema")
    add_text(
        doc,
        "A Tabela 10 consolida as métricas operacionais agregadas dos "
        "cinco agentes, capturadas em logs estruturados durante as "
        "execuções dos experimentos descritos nas Seções 6.1 a 6.5."
    )
    add_table(doc,
              ["Agente", "Latência p50", "Latência p95",
               "Tokens (médio)", "Custo (USD/inferência)"],
              [["Validador anti-PII", "0,02 ms", "0,03 ms", "n/a", "0,0000"],
               ["Narrador do PGR (com re-prompting)",
                "47,3 s", "72,5 s", "~3.500", "0,0012"],
               ["Gerador de Plano 5W2H",
                "7,1 s", "12,6 s", "~4.200", "0,0014"],
               ["Copiloto NR-01 (RAG)",
                "6,1 s", "8,8 s", "~1.800", "0,0004"],
               ["Analisador qualitativo",
                "~3 s", "~5 s", "~2.100", "0,0007"]],
              "Tabela 10 - Métricas operacionais por agente (gpt-4o-mini)",
              "Fonte: dados da pesquisa (2026), logs estruturados em "
              "ambiente local com chamadas reais à API OpenAI. Custo "
              "estimado para gpt-4o-mini (US$ 0,150/1M tokens entrada, "
              "US$ 0,600/1M tokens saída) e text-embedding-3-small "
              "(US$ 0,020/1M tokens), preços 2026.")
    add_text(
        doc,
        "A latência do validador anti-PII na ordem de dezenas de "
        "microssegundos confirma que o filtro pode operar como "
        "guard-rail obrigatório em todas as saídas da plataforma sem "
        "impacto sensível no tempo de resposta percebido pelo "
        "usuário. O tempo total da cadeia LLM é dominado pelas "
        "chamadas à API da OpenAI; o narrador é o agente com maior "
        "latência mediana (47,3 s) por incluir, em média, 1,3 "
        "tentativas adicionais de re-prompting para ajuste de "
        "tamanho. Em arquiteturas onde a interatividade é prioritária "
        "(copiloto e gerador 5W2H), os 6 a 10 segundos por resposta "
        "ficam dentro do envelope de aceitação tipicamente reportado "
        "pela literatura para QA médico assistido (XIONG et al., 2024)."
    )

    add_h2(doc, "6.7 Métricas de uso da plataforma")
    add_text(
        doc,
        "Em conformidade com as orientações do orientador, foram "
        "instrumentadas quatro métricas de uso operacional sobre o "
        "cenário de demonstração da plataforma (empresa fictícia com "
        "130 colaboradores em 8 setores, configurada via script "
        "determinístico de seed). A Tabela 11 consolida os tempos "
        "observados. Esses tempos correspondem ao cenário de "
        "demonstração agregada e dependem do tamanho do payload, "
        "da configuração de re-prompting e de cache de "
        "respostas equivalentes. Diferem, portanto, dos tempos por "
        "inferência reportados nas Tabelas 6, 7 e 10, que vêm de "
        "100 payloads de avaliação com re-prompting ativo e "
        "tamanhos variáveis."
    )
    add_table(doc,
              ["Métrica de uso", "Valor observado"],
              [["Tempo de processamento da campanha (130 respostas)",
                "1,8 s"],
               ["Tempo de devolução do Inventário pelo Narrador (payload demo, sem re-prompting)",
                "5,2 s (mediana)"],
               ["Tempo de geração do Plano de Ação 5W2H (cenário demo agregado)",
                "14,1 s (mediana)"],
               ["Tempo de exportação do PDF assinado",
                "0,9 s"],
               ["Taxa de sucesso na exportação PDF",
                "100% (n = 20)"]],
              "Tabela 11 - Métricas operacionais do site (cenário demo)",
              "Fonte: dados da pesquisa (2026), execução em ambiente "
              "local com seed determinístico. Latências menores que "
              "as Tabelas 6 e 10 porque o cenário demo usa payload "
              "menor e re-prompting desativado.")
    add_text(
        doc,
        "Os tempos confirmam que o ciclo completo de uma campanha, "
        "desde a tabulação dos resultados até a entrega de um Plano "
        "de Ação 5W2H auditável, ocorre em menos de 30 segundos. O "
        "gargalo é a chamada à API da OpenAI durante a geração do "
        "plano 5W2H, dominada pela latência da própria geração pelo "
        "LLM, não pela infraestrutura local."
    )


def build_discussao(doc):
    add_h1(doc, "7 DISCUSSÃO")
    add_text(
        doc,
        "Este capítulo articula os resultados apresentados no Capítulo "
        "6 com os trabalhos relacionados revisados no Capítulo 4, "
        "evidenciando ganhos, limitações e implicações da abordagem "
        "proposta."
    )

    add_h2(doc, "7.1 Anonimização em comparação com trabalhos correlatos")
    add_text(
        doc,
        "O validador anti-PII apresenta F1 "
        "doc-level de 0,9817 (IC95% [0,9722; 0,9901]) e Fβ=2 "
        "doc-level de 0,9710 (IC95% [0,9563; 0,9843]) sobre 1.025 "
        "casos sintéticos com anotação span-level. Em comparação "
        "com os dois trabalhos de referência em anonimização em "
        "português brasileiro: Schiezaro et al. (2026) reportam F1 "
        "de 0,9270 sobre 2.962 registros médicos clínicos reais, e "
        "Wiegand et al. (2024), no LLM-Anonymizer publicado em "
        "NEJM AI, reportam acurácia caractere a caractere de 0,9805 "
        "e recall de 0,9794 sobre aproximadamente 250 cartas "
        "clínicas em alemão e inglês. A comparação direta entre "
        "esses números é descritiva, não inferencial: os datasets "
        "diferem em natureza (sintético ocupacional vs. clínico real), "
        "tamanho (1.025 vs. 2.962 vs. 250), idioma (PT-BR vs. EN/DE) "
        "e granularidade da métrica reportada (F1 doc-level vs. F1 "
        "macro span-level vs. acurácia caractere a caractere)."
    )
    add_table(doc,
              ["Trabalho", "Idioma", "Técnica", "n",
               "Tipo de dado", "Métrica comparada",
               "Valor", "Latência"],
              [["Schiezaro et al. (2026)", "PT-BR",
                "BERT + LLM-as-judge", "2.962",
                "clínico real", "F1 doc-level",
                "0,9270", "segundos"],
               ["Wiegand et al. (2024)", "EN/DE",
                "LLM 70B local", "~250",
                "clínico real", "Acurácia caractere",
                "0,9805", "minutos"],
               ["Presente trabalho", "PT-BR",
                "regex + heurística", "1.025",
                "sintético ocupacional", "F1 doc-level / Fβ=2",
                "0,9817 / 0,9710", "microssegundos"]],
              "Tabela 13 - Comparação descritiva do validador anti-PII com trabalhos correlatos",
              "Fonte: elaborada pelos autores (2026). Atenção: a tabela compara "
              "ordens de grandeza, não desempenho equivalente em condições controladas.")
    add_text(
        doc,
        "A diferença de F1 entre o presente trabalho e Schiezaro et "
        "al. (2026) é +0,0547 a favor do primeiro, com a ressalva "
        "crítica de que a presente avaliação foi conduzida sobre "
        "amostras sintéticas controladas, enquanto Schiezaro et al. "
        "avaliaram dados clínicos reais com toda a heterogeneidade "
        "linguística associada. A vantagem real da arquitetura "
        "presente está no perfil de latência: o validador opera com "
        "p99 inferior a 0,1 ms (por design, regex e heurística vs. "
        "inferência neural), o que o torna adequado como guard-rail "
        "síncrono em pipelines de geração, sem custo perceptível "
        "para o usuário. Ainda assim, é possivelmente menos "
        "generalizável que abordagens baseadas em BERT ou LLM em "
        "dados não vistos, dado que a heurística depende de listas "
        "fechadas de prenomes e padrões regex; essa limitação está "
        "explícita na taxa de vazamento residual de 0,1455 para "
        "telefone e 0,0800 para matrícula e nome composto "
        "(Tabela 5b)."
    )
    add_text(
        doc,
        "Em síntese, a anonimização do presente trabalho atinge F1 "
        "comparável ao reportado por Schiezaro et al. (2026) em "
        "ordem de grandeza, com a ressalva de que a avaliação é em "
        "regime sintético e a generalização para clínico real "
        "permanece como hipótese a verificar."
    )

    add_h2(doc, "7.2 Geração de documentos técnicos")
    add_text(
        doc,
        "Voinea et al. (2024) reportaram, para a geração de laudos "
        "radiológicos, BERTScore F1 de 0,8054, ROUGE-1 F1 de 0,4998 "
        "e preferência humana em 21,8 por cento dos casos sobre o "
        "radiologista. O presente trabalho adota a mesma "
        "família de métricas (BERTScore F1 e ROUGE) sobre 100 "
        "payloads sintéticos, complementadas por faithfulness ao "
        "payload e taxa de alucinação regulatória, métricas de "
        "domínio específicas para o uso regulatório do narrador. "
        "O narrador atinge BERTScore F1 = 0,7623 (IC95% [0,7590; "
        "0,7660]), valor da mesma ordem de grandeza do 0,8054 de "
        "Voinea et al. (2024); contudo, os dois números não são "
        "comparáveis em sentido inferencial. Voinea et al. utilizaram "
        "como referência laudos radiológicos escritos por "
        "radiologistas humanos (ground truth de domínio independente "
        "do gerador), enquanto o inventário de referência usado aqui "
        "é gerado deterministicamente por template a partir do mesmo "
        "payload que alimenta o LLM. Esse compartilhamento de fonte "
        "tende a inflar a similaridade superficial, e a "
        "interpretação correta do número 0,7623 é: a saída do "
        "narrador preserva o vocabulário e a estrutura semântica "
        "esperados pelo template, não que ela equivale qualitativa "
        "ou clinicamente ao desempenho reportado por Voinea et al."
    )
    add_text(
        doc,
        "Faithfulness ao payload de 0,9316 é o número operacional "
        "central para o uso regulatório do narrador: aproximadamente "
        "93% das afirmações numéricas no Inventário gerado podem "
        "ser rastreadas a campos do payload de entrada, "
        "evidenciando ancoragem alta com cerca de 7% de afirmações "
        "sem suporte direto, tipicamente afirmações qualitativas "
        "(recomendações genéricas, contextualização técnica). "
        "Taxa de alucinação regulatória de 0,0000 sobre as "
        "citações de NR-01/NR-17/COPSOQ/ISO 45003 confirma que o "
        "narrador, no contexto deste experimento, não produz "
        "referências normativas inventadas. A conformidade na "
        "faixa de 600 a 900 palavras é de 0,6300 com o "
        "re-prompting iterativo ativado, evidenciando que a "
        "estratégia de re-prompting é elemento operacional "
        "obrigatório, não opcional, para alcançar a faixa de "
        "tamanho regulatório."
    )

    add_h2(doc, "7.3 Aderência à hierarquia de controle NR-01")
    add_text(
        doc,
        "A avaliação do gerador 5W2H quanto à aderência à "
        "hierarquia de controle NR-01 constitui métrica "
        "reproduzível para o recorte específico de aderência à "
        "hierarquia de controles para riscos psicossociais sob a "
        "NR-01, na ausência de baselines diretamente comparáveis "
        "na literatura consultada. A literatura sobre compliance "
        "automatizada (Hillebrand et al., 2025) trabalha com "
        "chatbots de orientação para outros regimes regulatórios, "
        "não com geração prescritiva para SST brasileira. A métrica "
        "operacional adotada é matriz de confusão 4x4 entre os "
        "níveis 1-Eliminação, 2-Substituição, 3-Controle "
        "Organizacional e 4-Controle Individual, conforme os "
        "incisos da Portaria MTE 1.419/2024."
    )
    add_text(
        doc,
        "O gold construído sem vazamento de rótulo "
        "(action_plan_gold_v3.py, 80 cenários cegos) e o prompt "
        "com árvore de decisão de quatro perguntas, definições "
        "essenciais por nível e seis exemplos few-shot resolvidos "
        "(três contrastando explicitamente Nível 1 versus Nível 2 "
        "e Nível 2 versus Nível 3) sustentam acurácia de 0,7875 "
        "(IC95% [0,7000; 0,8750]) e macro F1 de 0,7828. A matriz "
        "de confusão é equilibrada entre os quatro níveis: Nível 4 "
        "(Controle Individual) F1 = 0,952, Nível 3 (Controle "
        "Organizacional) F1 = 0,783, Nível 1 (Eliminação) F1 = "
        "0,765, e Nível 2 (Substituição) F1 = 0,632. O reforço "
        "contrastivo do prompt elevou o recall do Nível 2 a 0,600 "
        "(patamar prévio com prompt sem contraste explícito ficava "
        "em 0,350), evidenciando que a distinção fina entre "
        "substituir um processo perigoso e introduzir um controle "
        "organizacional sobre ele responde positivamente a "
        "engenharia de prompt orientada ao recorte, sem "
        "necessidade de fine-tuning. Os trabalhos futuros (Seção "
        "9.3) propõem aplicação Likert humana das rubricas "
        "specificity e implementability."
    )
    add_text(
        doc,
        "É importante separar duas avaliações distintas que se "
        "sobrepõem nesta seção. Primeiro, o pipeline de geração 5W2H "
        "(produção do JSON estruturado com os onze campos obrigatórios "
        "e as duas subseções do campo 'como') funciona: 100% dos "
        "planos gerados apresentaram estrutura completa, 91,25% "
        "referenciam ao menos um princípio andragógico de Knowles e "
        "0% apresentam alucinação regulatória em setenta e seis "
        "citações normativas. Segundo, a classificação na hierarquia "
        "NR-01 (escolha entre os quatro níveis de controle) atinge "
        "acurácia de 0,7875, com os quatro níveis apresentando F1 "
        "entre 0,632 e 0,952. "
        "Os dois resultados são complementares: o pipeline de "
        "geração preserva integridade estrutural e ancoragem "
        "regulatória, enquanto a classificação exige raciocínio "
        "normativo profundo sobre a natureza do risco. O conjunto "
        "demonstra que engenharia de prompt cuidadosa, combinada "
        "a um gold cego e bem rotulado, é suficiente para "
        "alcançar acurácia operacional próxima do alvo "
        "regulatório sem necessidade de fine-tuning."
    )

    add_h2(doc, "7.4 RAG sobre normas brasileiras vs. MIRAGE")
    add_text(
        doc,
        "Xiong et al. (2024) reportam, sobre 7.663 questões "
        "clínicas agregadas, ganho de 18 pontos percentuais em "
        "acurácia com RAG sobre Chain-of-Thought puro em LLMs "
        "comparáveis. O presente trabalho aplica RAG sobre um "
        "corpus normativo mais restrito (NR-01, NR-17, COPSOQ-II, "
        "ISO 45003 e LGPD) com gold de 40 perguntas. A execução "
        "com threshold cosseno 0,40 e corpus integral indexado "
        "(596 chunks de 600 caracteres cobrindo NR-01, NR-17 e "
        "LGPD nos textos completos, mais excertos curados de "
        "COPSOQ, ISO 45003 e política de k-anonymity) atinge "
        "Precision@1 = 0,6500 (IC95% [0,500; 0,800]), Recall@5 = "
        "0,8250 (IC95% [0,700; 0,925]) e MRR = 0,7271 "
        "(IC95% [0,6125; 0,8500]). O limiar de 0,40 é calibrado "
        "empiricamente para o modelo text-embedding-3-small em "
        "domínio normativo PT-BR, em que perguntas naturais "
        "raramente atingem cosseno superior a 0,75 com chunks "
        "técnicos."
    )
    add_text(
        doc,
        "O achado mais relevante da avaliação do copiloto é a "
        "aplicação do framework RAGAS (Es et al., 2024), que "
        "separa recuperação e geração. A faithfulness de 0,1514 "
        "(IC95% [0,0915; 0,2219]) revela que apenas 15% das "
        "sentenças geradas têm suporte semântico direto nos "
        "chunks recuperados (cosseno entre embedding da sentença "
        "e do melhor chunk acima de 0,75). Combinada com a "
        "context relevancy de 0,5508, essa métrica indica que o "
        "LLM responde predominantemente com seu conhecimento "
        "paramétrico em vez de interpretar os chunks recuperados, "
        "fenômeno que Es et al. denominam 'shortcut generation' "
        "em domínios onde o LLM já foi exposto à norma durante "
        "pré-treinamento. A answer relevancy de 0,6116 confirma "
        "que as respostas, mesmo sem ancoragem, endereçam o "
        "tópico da pergunta. Em contraste, a taxa de alucinação "
        "regulatória de 0,0000 sobre dezenove citações é o sinal "
        "operacional mais positivo do agente: nenhuma referência "
        "normativa produzida aponta para item inexistente. O "
        "resultado decorre do prompt com regra de ouro de "
        "abstenção, da indexação dos textos integrais de NR-01, "
        "NR-17 e LGPD (596 chunks indexados versus excertos "
        "anteriormente) e do lookup hierárquico no "
        "regulatory_lookup.py, que aceita subitens detalhados "
        "como prefixos válidos de itens da base curada. A "
        "faithfulness baixa continua motivando trabalhos futuros "
        "na Seção 9.2 (reranking pós-retrieval via cross-encoder "
        "e fine-tuning de instrução para ancoragem)."
    )

    add_h2(doc, "7.5 Clustering temático")
    add_text(
        doc,
        "A literatura recente (Jiang, Liu e Fisher, 2026) reporta "
        "que BERTopic atinge maior coerência tópica c_v do que "
        "Structural Topic Models (STM) em respostas curtas de "
        "pesquisas de opinião genéricas, enquanto STM mantém "
        "vantagem em análise inferencial com covariáveis. O "
        "presente trabalho adota abordagem alternativa, delegando "
        "o clustering a um LLM com prompt estruturado, sob duas "
        "vantagens declaradas: controle explícito do número máximo "
        "de clusters e da nomenclatura em português, e descrição "
        "semântica do cluster gerada na mesma chamada, eliminando "
        "a etapa de rotulação humana."
    )
    add_text(
        doc,
        "A presente avaliação compara empiricamente o LLM contra os três "
        "baselines previstos: k-means k=7 sobre embeddings "
        "text-embedding-3-small, LDA com sete tópicos e BERTopic "
        "com UMAP+HDBSCAN. Os resultados (Tabela 9) revelam padrão "
        "duplo. Em concordância categorial (ARI), o LLM atinge "
        "0,1498 (desvio-padrão 0,0069 sobre cinco seeds), abaixo "
        "do k-means (0,3681) e do LDA (0,1625), mas acima do "
        "BERTopic (0,1133). Em coerência tópica c_v, métrica "
        "primária de topic modeling, o LLM lidera com 0,4724 "
        "(desvio-padrão 0,1057), contra 0,3517 do k-means, 0,3552 "
        "do BERTopic e 0,3460 do LDA. LDA é prejudicado pela "
        "natureza bag-of-words que ignora sinônimos em respostas "
        "curtas; BERTopic é prejudicado pelo tamanho amostral "
        "pequeno (n = 150 abaixo do regime favorável a "
        "UMAP+HDBSCAN com hiperparâmetros padrão). A dissociação "
        "ARI baixo / c_v alto no LLM é coerente com a observação "
        "de Röder, Both e Hinneburg (2015): clusters "
        "semanticamente coesos podem não coincidir com a "
        "categorização do gold."
    )
    add_text(
        doc,
        "À luz desse padrão duplo, a escolha do LLM como método "
        "principal mantém-se justificada pelas vantagens "
        "operacionais (nomes descritivos por cluster sem etapa "
        "adicional de rotulação, controle explícito do número "
        "máximo de clusters via prompt e descrição semântica do "
        "tema na mesma chamada) e pela liderança em coerência "
        "tópica c_v, com perda apenas em ARI frente ao k-means. "
        "Os trabalhos futuros descritos no "
        "Capítulo 9 propõem três caminhos: (i) prompt menos "
        "restritivo que permita o LLM redefinir suas fronteiras "
        "temáticas dentro de um limite máximo de clusters; (ii) "
        "ensemble LLM+k-means com voto majoritário; (iii) gold "
        "ampliado para 500 respostas e definição operacional "
        "dos temas Carreira e Outros, que continuam sendo o "
        "ponto de colapso."
    )

    add_h2(doc, "7.6 Limitações observadas e ameaças à validade")
    add_text(
        doc,
        "A análise crítica dos resultados, organizada segundo o "
        "esquema de ameaças à validade de Wohlin et al. (2012), "
        "identifica quatro dimensões. Validade de construto: as "
        "métricas escolhidas (Fβ=2, BERTScore, faithfulness, c_v) "
        "capturam parcialmente as propriedades de interesse; "
        "BERTScore favorece paráfrase semântica mas não detecta "
        "erro factual sutil; faithfulness por embedding similarity "
        "depende de threshold arbitrário; c_v depende do tamanho de "
        "janela. Validade interna: o gold standard do 5W2H, na "
        "versão original (n=20), continha vazamento de rótulo via "
        "descrição da natureza do risco, invalidando o F1 de 1,000 "
        "como evidência; o gold v3 (n=80) elimina o vazamento "
        "lexical mas reconhece resíduo inerente: a natureza do "
        "risco psicossocial está intrinsecamente ligada ao nível "
        "apropriado da hierarquia de controle."
    )
    add_text(
        doc,
        "Validade externa: toda a avaliação foi conduzida sobre "
        "dados sintéticos algoritmicamente gerados; a "
        "generalização para dados clínicos ou ocupacionais reais "
        "depende de fatores que não estão modelados nos cenários "
        "sintéticos (heterogeneidade linguística, ortografia "
        "irregular, marcadores regionais de PII), e a validação em "
        "campo, sob aprovação ética, permanece como trabalho futuro "
        "inegociável. Os oito setores cobertos não esgotam a "
        "heterogeneidade ocupacional do Brasil. Validade de "
        "conclusão: o tamanho amostral por agente (n=1.025 para "
        "anti-PII, n=80 para 5W2H, n=100 para narrador, "
        "n=40 para copiloto, n=150 para "
        "analisador qualitativo) é compatível com bootstrap "
        "não-paramétrico de 1.000 reamostragens; comparações "
        "pareadas dentro de um mesmo dataset são inferenciais via "
        "McNemar; comparações entre datasets distintos (literatura) "
        "permanecem descritivas."
    )
    add_text(
        doc,
        "Outras limitações operacionais: (i) falsos negativos "
        "residuais em fronteiras de regex (alguns formatos "
        "internacionais de telefone, nomes próprios sem acento "
        "e matrículas sem qualquer gatilho textual) "
        "permanecem; (ii) a avaliação humana Likert, "
        "embora prevista com dois revisores e kappa de Cohen "
        "ponderado, teve rubricas e amostras preparadas (rubricas pré-registradas e amostras "
        "exportadas) mas a aplicação ainda não foi concluída; (iii) "
        "dependência de API externa OpenAI, mitigada por "
        "arquitetura com drop-in replacement (LangChain) que "
        "permite substituição por modelos locais."
    )

    add_h2(doc, "7.7 Implicações práticas e científicas")
    add_text(
        doc,
        "A combinação avaliada (validador anti-PII determinístico "
        "mais agentes LLM com saída estruturada mais piso de "
        "k-anonymity) demonstra, em regime sintético, ser possível "
        "construir uma cadeia de geração de documentos regulatórios "
        "em saúde ocupacional com rastreabilidade, anonimato e "
        "custo operacional acessível. Cientificamente, o trabalho "
        "oferece três contribuições reproduzíveis: o gold standard "
        "ampliado de 1.025 casos com anotação span-level para PII "
        "em PT-BR ocupacional; o protocolo de avaliação "
        "com métricas adequadas à função de utilidade de cada "
        "agente (Fβ=2, BERTScore, RAGAS, c_v) e intervalos "
        "de confiança bootstrap; e a arquitetura de guard-rail "
        "síncrono que reconcilia LLM e exigências regulatórias. "
        "Todos os artefatos (backend, frontend, scripts de avaliação, "
        "gold standards e gerador determinístico do documento) estão "
        "publicamente disponíveis em "
        "https://github.com/Alencar-png/emotion-care-tcc, sob o "
        "mesmo controle de versão usado durante a pesquisa. A "
        "transferência para uso em produção requer, contudo, a "
        "validação em campo descrita na Seção 7.6."
    )

    add_h2(doc, "7.8 Comparação com LLMs open-source locais")
    add_text(
        doc,
        "Para investigar se a dependência da API proprietária da "
        "OpenAI é tecnicamente necessária ou apenas conveniente, "
        "esta pesquisa executou o mesmo protocolo de avaliação dos "
        "quatro agentes generativos (narrador, gerador 5W2H, "
        "copiloto RAG e analisador qualitativo) sobre dois modelos "
        "open-source de 14 bilhões de parâmetros executados "
        "localmente em GPU AMD RX 9060 XT (16 GB VRAM, backend "
        "Vulkan) via Ollama: Qwen 2.5 14B (ALIBABA, 2024) e Phi-4 "
        "14B (MICROSOFT, 2024). Os gold standards, as métricas e "
        "os critérios de bootstrap permanecem idênticos aos "
        "descritos no Capítulo 5; apenas o backend de geração "
        "muda. A Tabela 12 consolida os resultados primários."
    )
    add_table(doc,
              ["Agente / Métrica", "gpt-4o-mini", "Qwen 2.5 14B",
               "Phi-4 14B"],
              [["5W2H - Acurácia 4 níveis", "0,7875", "0,7375",
                "0,7375"],
               ["5W2H - Macro F1", "0,7828", "0,7363", "0,7158"],
               ["5W2H - Andragogia (Knowles)", "0,9250", "0,8625",
                "0,7750"],
               ["5W2H - Latência p50 (s)", "7,1", "14,1", "14,9"],
               ["Copilot - Precision@1", "0,6500", "0,8750",
                "0,8000"],
               ["Copilot - MRR", "0,7271", "0,9250", "0,8708"],
               ["Copilot - Faithfulness (RAGAS)", "0,1514",
                "0,2193", "0,1614"],
               ["Copilot - Alucinação regulatória", "0,0000",
                "0,2500", "0,0000"],
               ["Copilot - Latência p50 (s)", "5,8", "11,2",
                "28,7"],
               ["Narrador - BERTScore F1", "0,7623", "0,7655",
                "0,7680"],
               ["Narrador - Faithfulness payload", "0,9316",
                "0,8944", "0,8836"],
               ["Narrador - Word range 600-900", "0,6300",
                "0,5600", "0,9400"],
               ["Narrador - Latência p50 (s)", "47,3", "124,2",
                "117,8"],
               ["Qualitativo - ARI (5 seeds)", "0,1498", "0,1423",
                "0,1423"],
               ["Qualitativo - c_v (5 seeds)", "0,4724", "0,3565",
                "0,3565"]],
              "Tabela 12 - Comparação dos quatro agentes generativos "
              "entre gpt-4o-mini (API OpenAI) e LLMs open-source 14B "
              "rodando localmente",
              "Fonte: dados da pesquisa (2026). Ollama 0.4 com "
              "Vulkan ativo (AMD RX 9060 XT 16GB).")

    add_figure(doc, FIG / "fig_compare_action_plan.png",
               "Figura 8 - Gerador 5W2H, F1 por nível NR-01 (3 modelos)",
               "Fonte: dados da pesquisa (2026).")
    add_figure(doc, FIG / "fig_compare_copilot_ragas.png",
               "Figura 9 - Copiloto NR-01, métricas RAGAS (3 modelos)",
               "Fonte: dados da pesquisa (2026).")
    add_figure(doc, FIG / "fig_compare_narrator.png",
               "Figura 10 - Narrador do PGR, métricas primárias (3 modelos)",
               "Fonte: dados da pesquisa (2026).")
    add_figure(doc, FIG / "fig_compare_qualitative.png",
               "Figura 11 - Analisador qualitativo, ARI/NMI/c_v/F1 (3 modelos, 5 seeds)",
               "Fonte: dados da pesquisa (2026).")
    add_figure(doc, FIG / "fig_pareto_latencia_qualidade.png",
               "Figura 12 - Pareto latência mediana versus qualidade "
               "(média geométrica) dos três modelos",
               "Fonte: dados da pesquisa (2026).")
    add_figure(doc, FIG / "fig_antipii_v20_v21.png",
               "Figura 13 - Anti-PII Fβ=2 por categoria, evolução do "
               "validador determinístico entre versões",
               "Fonte: dados da pesquisa (2026).")

    add_text(
        doc,
        "Os achados centrais da comparação são quatro. "
        "Primeiro, o gpt-4o-mini lidera no gerador de Plano 5W2H, "
        "tanto em acurácia (0,7875 versus 0,7375 nos dois "
        "open-source) quanto em aderência andragógica de Knowles "
        "(0,9250 versus 0,8625 do Qwen e 0,7750 do Phi-4), com "
        "F1 do Nível 2 (Substituição) elevado a 0,632 pelo prompt "
        "contrastivo, superando o Phi-4 (0,6000) e Qwen (0,5854) "
        "em todos os quatro níveis."
    )
    add_text(
        doc,
        "Segundo, o copiloto NR-01 com RAGAS mostra padrão "
        "interessante: o Qwen 2.5 14B atinge a maior Precision@1 "
        "(0,8750) e MRR (0,9250), seguido pelo Phi-4 14B "
        "(Precision@1 = 0,8000; MRR = 0,8708); o gpt-4o-mini, com "
        "corpus integral e prompt revisado, atinge Precision@1 = "
        "0,6500 e MRR = 0,7271, valores intermediários. Em "
        "alucinação regulatória, gpt-4o-mini e Phi-4 14B empatam "
        "no patamar ideal de 0,0000, enquanto o Qwen ainda produz "
        "1 citação inválida em cada 4. Esse padrão sugere que o "
        "gpt-4o-mini e o Phi-4, sob a combinação adequada de "
        "prompt de abstenção e base de validação hierárquica, são "
        "estruturalmente preferíveis em ambientes regulatórios "
        "onde alucinação zero é requisito hard."
    )
    add_text(
        doc,
        "Terceiro, no narrador do PGR, o Phi-4 14B obtém o "
        "maior word range na faixa 600 a 900 palavras (0,9400, "
        "contra 0,6300 do gpt-4o-mini e 0,5600 do Qwen), o que "
        "indica conformidade estrutural superior ao modelo "
        "proprietário; em contrapartida, a faithfulness ao "
        "payload (ancoragem numérica) é maior no gpt-4o-mini "
        "(0,9316 versus 0,8836 do Phi-4 e 0,8944 do Qwen). "
        "BERTScore F1 fica praticamente equivalente entre os "
        "três (0,7623 a 0,7680), reforçando que a similaridade "
        "semântica global ao inventário de referência é "
        "característica robusta da família LLM 14B+ e não "
        "exclusiva do gpt-4o-mini."
    )
    add_text(
        doc,
        "Quarto, no analisador qualitativo os três modelos "
        "produzem resultados muito próximos no gpt-4o-mini, Qwen "
        "e Phi-4 (ARI entre 0,142 e 0,150 a temperatura zero), "
        "confirmando que o gargalo desse agente é o prompt "
        "enrijecido com fronteiras semânticas e o tamanho do "
        "gold, não a capacidade do modelo. O LLM lidera, contudo, "
        "na coerência tópica c_v (0,4724 do gpt-4o-mini contra "
        "0,3517 do k-means), métrica primária de topic modeling."
    )
    add_text(
        doc,
        "A análise Pareto (Figura 12) sintetiza o trade-off "
        "latência-qualidade. O gpt-4o-mini é dominante na "
        "fronteira inferior do eixo x (latência mediana 17,2 s "
        "média entre os quatro agentes), enquanto o Qwen 2.5 14B "
        "domina a fronteira superior do eixo y (qualidade média "
        "geométrica 0,5031). O Phi-4 14B é Pareto-eficiente em "
        "cenários onde alucinação zero é requisito hard (copilot "
        "regulatório). Em termos de custo, os dois open-source "
        "apresentam custo marginal próximo de zero por inferência "
        "(apenas energia elétrica local), contra US$ 0,0004 a "
        "US$ 0,0014 por inferência do gpt-4o-mini. Para "
        "instituições do tipo CESMAC/UNIMA que pretendam adotar "
        "o sistema em escala (centenas de campanhas/ano), o "
        "custo total de propriedade favorece a substituição do "
        "gpt-4o-mini por Phi-4 ou Qwen em agentes específicos "
        "(copiloto e narrador), preservando o gpt-4o-mini apenas "
        "onde a vantagem qualitativa é robusta (gerador 5W2H)."
    )


def build_conclusoes(doc):
    add_h1(doc, "8 CONCLUSÕES")
    add_text(
        doc,
        "Este trabalho investigou a viabilidade do uso de Modelos "
        "de Linguagem de Grande Escala combinados a Geração "
        "Aumentada por Recuperação como mecanismo de interpretação "
        "de respostas do COPSOQ II-Br e geração automatizada dos "
        "documentos exigidos pelo Programa de Gerenciamento de "
        "Riscos da NR-01. A investigação foi conduzida sobre uma "
        "plataforma experimental (Emotion Care) que serviu de "
        "ambiente controlado para concepção, instrumentação e "
        "avaliação de cinco agentes especializados."
    )
    add_text(
        doc,
        "Os resultados preliminares sobre dados sintéticos sugerem "
        "viabilidade técnica da arquitetura combinada e indicam a "
        "próxima etapa: validação com dados reais sob aprovação de "
        "comitê de ética em pesquisa. Em particular, o validador "
        "anti-PII apresentou Fβ=2 doc-level de 0,9710 "
        "(IC95% [0,9563; 0,9843]) e Fβ=2 macro span-level de "
        "0,9688 (IC95% [0,9546; 0,9822]) sobre 1.025 casos rotulados "
        "com anotação span-level, com latência de pico abaixo de "
        "0,1 ms. O número atinge F1 doc-level comparável ao "
        "reportado por Schiezaro et al. (2026) em ordem de "
        "grandeza (0,9817 vs. 0,9270), com a ressalva crítica de "
        "que a presente avaliação foi conduzida sobre amostras "
        "sintéticas controladas, enquanto Schiezaro et al. "
        "avaliaram 2.962 registros clínicos reais. Comparações "
        "diretas devem ser tomadas como referência de ordem de "
        "grandeza, não como superação metodológica."
    )
    add_text(
        doc,
        "Os agentes baseados em LLM (narrador do PGR, gerador 5W2H, "
        "copiloto NR-01 e analisador qualitativo) foram avaliados "
        "sob o aparato metodológico descrito na Seção 5.6. O "
        "narrador atingiu BERTScore F1 = 0,7623 "
        "(IC95% [0,7590; 0,7660]) e faithfulness ao payload de "
        "0,9316 (IC95% [0,9153; 0,9466]), com taxa de alucinação "
        "regulatória de 0,0000 sobre citações de "
        "NR-01/NR-17/COPSOQ/ISO 45003. O copiloto, sob framework "
        "RAGAS (Es et al., 2024), apresentou Precision@1 = 0,6500, "
        "Recall@5 = 0,8250 (IC95% [0,7000; 0,9250]), MRR = 0,7271, "
        "faithfulness de 0,1514 (IC95% [0,0915; 0,2219]) e taxa de "
        "alucinação regulatória de 0,0000 sobre as citações "
        "produzidas, sobre corpus integral indexado de NR-01, "
        "NR-17 e LGPD (596 chunks). A faithfulness modesta "
        "configura limitação central do trabalho: o componente "
        "RAG não demonstra ancoragem efetiva, levando o LLM a "
        "responder com conhecimento paramétrico em vez de "
        "interpretar os chunks recuperados. A correção desse "
        "comportamento via reranking pós-retrieval e fine-tuning "
        "de instrução está detalhada no Capítulo 9. "
        "O analisador qualitativo, comparado contra três baselines "
        "(BERTopic, k-means k=7 sobre embeddings e LDA), atingiu "
        "ARI = 0,1498 ± 0,0069 (média de cinco seeds, "
        "temperatura zero), abaixo de k-means (0,3681) em "
        "concordância categorial, mas com coerência tópica c_v "
        "de 0,4724, superior aos três baselines (k-means 0,3517, "
        "BERTopic 0,3552, LDA 0,3460), o que sustenta o uso do "
        "LLM como método principal apesar da perda em ARI."
    )
    add_text(
        doc,
        "O gerador 5W2H, sob o gold cego sem vazamento de rótulo "
        "(80 cenários, action_plan_gold_v3.py), atinge acurácia "
        "de 0,7875 (IC95% [0,7000; 0,8750]) na hierarquia de "
        "quatro níveis da NR-01, sustentada pela árvore de "
        "decisão e pelos seis exemplos few-shot do prompt, com "
        "três casos contrastando explicitamente Nível 1 versus "
        "Nível 2 e Nível 2 versus Nível 3. A matriz de confusão "
        "é equilibrada nos quatro níveis: Nível 4 (Controle "
        "Individual) F1 = 0,952, Nível 3 (Controle "
        "Organizacional) F1 = 0,783, Nível 1 (Eliminação) F1 = "
        "0,765, Nível 2 (Substituição) F1 = 0,632. A aderência "
        "andragógica é de 0,9250 e a taxa de alucinação "
        "regulatória é de 0,0000. Os trabalhos futuros descritos "
        "no Capítulo 9 incluem aplicação Likert humana das "
        "rubricas specificity e implementability pelos dois "
        "revisores independentes."
    )
    add_text(
        doc,
        "Este trabalho contribui ainda com um benchmark inédito "
        "de comparação tripla entre o gpt-4o-mini (OpenAI API) "
        "e dois LLMs open-source de 14 bilhões de parâmetros "
        "executados localmente (Qwen 2.5 14B e Phi-4 14B via "
        "Ollama em GPU AMD RX 9060 XT), documentado na Seção "
        "7.8. Os achados-chave: o Qwen supera o gpt-4o-mini em "
        "Precision@1 no copiloto NR-01 (0,8750 contra 0,6500) e "
        "em MRR (0,9250 contra 0,7271); gpt-4o-mini e Phi-4 "
        "atingem alucinação regulatória de exatamente 0,0000 no "
        "copiloto; o Phi-4 14B atinge ainda word range estrutural "
        "de 0,9400 no narrador, superior ao gpt-4o-mini; e o "
        "gpt-4o-mini preserva vantagem no gerador 5W2H (acurácia "
        "0,7875 vs. 0,7375 dos open-source) e na faithfulness ao "
        "payload do narrador. Esse benchmark sustenta a viabilidade "
        "técnica e econômica de migração parcial para LLMs "
        "open-source, especialmente nos agentes regulatórios "
        "onde alucinação zero é requisito hard, preservando "
        "soberania de dados e eliminando custo marginal por "
        "inferência."
    )
    add_text(
        doc,
        "Sob a perspectiva metodológica, três princípios foram "
        "deliberadamente adotados para evitar inflação artificial "
        "dos resultados: gold cego no gerador 5W2H, sem pistas "
        "lexicais que sinalizem ao modelo a natureza do controle "
        "esperado; substituição de F1 por Fβ=2 no anti-PII, "
        "métrica adequada à assimetria entre falso positivo e "
        "falso negativo sob LGPD; e abandono de AUC/ROC para "
        "classificador determinístico sem score contínuo. Esse "
        "rigor metodológico, aliado a bootstrap não-paramétrico "
        "para todos os intervalos de confiança, sustenta que os "
        "números apresentados refletem desempenho real, "
        "verificável e reproduzível a partir do repositório "
        "público."
    )
    add_text(
        doc,
        "Sob a perspectiva da Ciência da Computação, o trabalho "
        "oferece três contribuições reproduzíveis: o gold standard "
        "ampliado de 1.025 casos com anotação span-level para PII "
        "em PT-BR ocupacional, ampliável a partir do gerador "
        "determinístico publicado; o protocolo de avaliação "
        "definido com métricas justificadas pela função de "
        "utilidade de cada agente (Fβ=2 para anti-PII, BERTScore "
        "e faithfulness e taxa de alucinação regulatória para "
        "agentes generativos, framework RAGAS para o copiloto, "
        "coerência c_v para o analisador qualitativo); e a "
        "arquitetura de guard-rail síncrono que reconcilia LLM e "
        "exigências regulatórias. O código-fonte completo dos cinco "
        "agentes, do frontend, dos scripts de avaliação, dos gold "
        "standards e do gerador determinístico do TCC está "
        "publicamente disponível em "
        "https://github.com/Alencar-png/emotion-care-tcc para "
        "verificação e extensão por outros pesquisadores. A "
        "pergunta-problema formulada na Introdução é respondida "
        "afirmativamente em regime de dados sintéticos, restando "
        "à validação com dados reais a "
        "confirmação em campo."
    )


def build_trabalhos_futuros(doc):
    add_h1(doc, "9 SUGESTÕES PARA TRABALHOS FUTUROS")
    add_text(
        doc,
        "Os trabalhos futuros estão organizados em três grupos de "
        "prioridade decrescente. As quatro primeiras frentes (9.1 a "
        "9.4) endereçam diretamente limitações centrais identificadas "
        "no Capítulo 7 e são pré-requisitos para uso em produção. "
        "As demais (9.5) descrevem extensões de escopo desejáveis, "
        "porém não bloqueantes."
    )
    add_h2(doc, "9.1 Validação com dados reais sob aprovação ética")
    add_text(
        doc,
        "Conduzir avaliação dos cinco agentes em pelo menos duas "
        "organizações reais, com respondentes humanos, mediante "
        "aprovação do Comitê de Ética em Pesquisa. Pré-requisito "
        "inegociável: toda a avaliação reportada neste TCC foi "
        "conduzida sobre dados sintéticos, e a generalização para "
        "campo permanece como hipótese a verificar (Seção 7.6)."
    )
    add_h2(doc, "9.2 Reranking pós-retrieval e fine-tuning de ancoragem")
    add_text(
        doc,
        "Introduzir reranker pós-retrieval baseado em cross-encoder "
        "(por exemplo, bge-reranker-v2-m3 ou cohere-rerank-3) sobre "
        "os top-20 candidatos do retriever cosseno e selecionar os "
        "top-5 reranqueados antes de passar ao gerador. Em paralelo, "
        "explorar fine-tuning de instrução do gpt-4o-mini com pares "
        "(contexto, pergunta, resposta-ancorada) para reforçar a "
        "ancoragem semântica das sentenças geradas. "
        "Justificativa: a faithfulness de 0,1514 observada na "
        "Seção 6.4 demonstra que o LLM responde com conhecimento "
        "paramétrico em vez de interpretar os chunks recuperados, "
        "mesmo com prompt rigoroso de abstenção e corpus integral "
        "indexado. Meta: faithfulness mínima de 0,60, mantendo a "
        "alucinação regulatória abaixo de 0,05."
    )
    add_h2(doc, "9.3 Fine-tuning para o Nível 2 da hierarquia 5W2H")
    add_text(
        doc,
        "O gerador 5W2H atinge acurácia global de 0,7875 e F1 = "
        "0,952 no Nível 4 (Controle Individual), com o Nível 2 "
        "(Substituição) ainda no patamar mais baixo, F1 = 0,632 "
        "(Seção 6.3). Conduzir fine-tuning supervisionado do "
        "gpt-4o-mini com pares cenário-rótulo balanceados nos "
        "quatro níveis, com ênfase em discriminação fina "
        "Substituição versus Controle Organizacional. Meta: F1 "
        "mínimo de 0,80 em todos os quatro níveis."
    )
    add_h2(doc, "9.4 Aplicação Likert humana e medição de kappa")
    add_text(
        doc,
        "Concluir a avaliação Likert pré-registrada das dimensões "
        "specificity e implementability do 5W2H e clareza percebida "
        "do narrador e do copiloto, com os dois revisores aplicando "
        "as rubricas pré-registradas (módulo likert_rubrics.py) sobre "
        "as amostras estratificadas já exportadas. Reportar kappa de "
        "Cohen ponderado e resolução de divergências para "
        "diferenças maiores ou iguais a dois pontos."
    )
    add_h2(doc, "9.5 Extensões de escopo")
    add_text(
        doc,
        "Direções adicionais, ordenadas por afinidade temática: "
        "(i) avaliação com modelos open-source locais (Llama-3, "
        "Mistral, BERTimbau, PTT5-v2) para reduzir dependência de "
        "API comercial; (ii) re-execução do narrador com "
        "re-prompting iterativo ativado para reportar conformidade "
        "de tamanho pós-ajuste; (iii) expansão do catálogo de "
        "instrumentos psicométricos (MBI, QPS Nordic, EET, EACT); "
        "(iv) integração com eSocial e assinatura digital ICP-Brasil "
        "para os documentos gerados; e (v) avaliação de usabilidade "
        "conforme ISO 9241-11 e auditoria de segurança independente "
        "conforme metodologias OWASP."
    )


def build_referencias(doc):
    add_h1(doc, "REFERÊNCIAS")
    refs = [
        "ALIBABA CLOUD. Qwen2.5 technical report. arXiv preprint arXiv:2412.15115, 2024.",
        "ALSAAD, R.; ALSHAKHS, S.; THOMAS, R. Depression subtype classification from social media posts: few-shot prompting vs. fine-tuning of large language models. Frontiers in Digital Health, [s.l.], 2026.",
        "BORATTI, R. R.; ROCHA, K. B.; SANTOS, F. T. Adaptação transcultural do Copenhagen Psychosocial Questionnaire (COPSOQ-II) para o português brasileiro. Revista Brasileira de Saúde Ocupacional, v. 43, p. 1-12, 2018.",
        "BRASIL. Lei n. 13.709, de 14 de agosto de 2018. Lei Geral de Proteção de Dados Pessoais (LGPD). Diário Oficial da União, Brasília, 15 ago. 2018.",
        "BRASIL. Ministério do Trabalho e Emprego. Norma Regulamentadora NR-01: Disposições Gerais e Gerenciamento de Riscos Ocupacionais. Brasília: MTE, 2022.",
        "BRASIL. Ministério do Trabalho e Emprego. Norma Regulamentadora NR-17: Ergonomia. Brasília: MTE, 2018.",
        "BRASIL. Ministério do Trabalho e Emprego. Portaria nº 1.419, de 27 de agosto de 2024. Altera a Norma Regulamentadora nº 01 (NR-01) para incluir os fatores de risco psicossocial relacionados ao trabalho. Diário Oficial da União, Brasília, 28 ago. 2024.",
        "BRASIL. Ministério do Trabalho e Emprego. Portaria nº 765, de 15 de maio de 2025. Estabelece o prazo de adequação para fiscalização das disposições da NR-01 relativas aos fatores de risco psicossocial, fixando o início da vigência fiscalizável em 26 de maio de 2026. Diário Oficial da União, Brasília, 16 maio 2025.",
        "COX, T.; GRIFFITHS, A. The nature and measurement of work-related psychosocial hazards: a frame of reference for the assessment of work stressors. Work & Stress, v. 9, n. 3-4, p. 244-261, 1995.",
        "EFRON, B.; TIBSHIRANI, R. J. An Introduction to the Bootstrap. New York: Chapman & Hall, 1993.",
        "ES, S.; JAMES, J.; ESPINOSA-ANKE, L.; SCHOCKAERT, S. RAGAS: Automated Evaluation of Retrieval Augmented Generation. In: Proceedings of the 18th Conference of the European Chapter of the Association for Computational Linguistics (EACL 2024): System Demonstrations. Stroudsburg: ACL, 2024.",
        "FEW, S. Now You See It: simple visualization techniques for quantitative analysis. Burlingame: Analytics Press, 2009.",
        "GIL, A. C. Como elaborar projetos de pesquisa. 4. ed. São Paulo: Atlas, 2002.",
        "HILLEBRAND, L.; BERGER, A.; UEDELHOVEN, D.; BERGHAUS, D.; WARNING, U.; DILMAGHANI, T.; KLIEM, B.; SCHMID, T.; LOITZ, R.; SIFA, R. Advancing Risk and Quality Assurance: a RAG Chatbot for Improved Regulatory Compliance. arXiv preprint arXiv:2507.16711, 2025.",
        "INSTITUTO NACIONAL DO SEGURO SOCIAL (INSS). Anuário Estatístico Previdenciário 2023. Brasília: INSS, 2024.",
        "INTERNATIONAL ORGANIZATION FOR STANDARDIZATION. ISO 45003: Occupational health and safety management. Psychological health and safety at work. Guidelines for managing psychosocial risks. Geneva: ISO, 2021.",
        "JIANG, Y.; LIU, S.; FISHER, P. A. A comparative evaluation of structural topic models and BERTopic for short, open-ended survey responses. arXiv preprint arXiv:2605.23093, 2026.",
        "KARASEK, R.; THEORELL, T. Healthy Work: stress, productivity, and the reconstruction of working life. New York: Basic Books, 1990.",
        "KNOWLES, M. S.; HOLTON III, E. F.; SWANSON, R. A. The Adult Learner: the definitive classic in adult education and human resource development. 8. ed. London: Routledge, 2015.",
        "KRISTENSEN, T. S.; HANNERZ, H.; HOGH, A.; BORG, V. The Copenhagen Psychosocial Questionnaire: a tool for the assessment and improvement of the psychosocial work environment. Scandinavian Journal of Work, Environment & Health, v. 31, n. 6, p. 438-449, 2010.",
        "LEWIS, P.; PEREZ, E.; PIKTUS, A.; PETRONI, F.; KARPUKHIN, V. et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. Advances in Neural Information Processing Systems (NeurIPS), v. 33, p. 9459-9474, 2020.",
        "MANNING, C. D.; RAGHAVAN, P.; SCHÜTZE, H. Introduction to Information Retrieval. Cambridge: Cambridge University Press, 2008.",
        "MASLACH, C.; JACKSON, S. E. The measurement of experienced burnout. Journal of Organizational Behavior, v. 2, n. 2, p. 99-113, 1981.",
        "MICROSOFT. Phi-4 technical report. arXiv preprint arXiv:2412.08905, 2024.",
        "ORGANIZAÇÃO INTERNACIONAL DO TRABALHO (OIT). Mental Health at Work: policy brief. Genebra: ILO, 2022.",
        "ORGANIZAÇÃO MUNDIAL DA SAÚDE (OMS). WHO Guidelines on Mental Health at Work. Genebra: WHO, 2022.",
        "RÖDER, M.; BOTH, A.; HINNEBURG, A. Exploring the space of topic coherence measures. In: WSDM '15: Proceedings of the Eighth ACM International Conference on Web Search and Data Mining. Shanghai: ACM, 2015. p. 399-408.",
        "SCHIEZARO, M.; ROSA, G.; CAMPOS, B. A. G.; PEDRINI, H. Guardians of the data: NER and LLMs for effective medical record anonymization in Brazilian Portuguese. Frontiers in Public Health, [s.l.], 2026.",
        "SHAO, Z.; WANG, X.; LIU, Z.; WANG, C.; SUBBALAKSHMI, K. P. Systematic evaluation of machine-generated reasoning and PHQ-9 labeling for depression detection using large language models. arXiv preprint arXiv:2505.17119, 2025.",
        "SHI, H.; LIU, J.; YANG, C.; SHANG, J.; ZENG, Y. Machine learning for predicting burnout among healthcare workers: a systematic review and meta-analysis. Contemporary Nurse, [s.l.], 2025.",
        "SIEGRIST, J. Adverse health effects of high-effort/low-reward conditions. Journal of Occupational Health Psychology, v. 1, n. 1, p. 27-41, 1996.",
        "SOKOLOVA, M.; LAPALME, G. A systematic analysis of performance measures for classification tasks. Information Processing & Management, v. 45, n. 4, p. 427-437, 2009.",
        "SWEENEY, L. k-anonymity: a model for protecting privacy. International Journal on Uncertainty, Fuzziness and Knowledge-Based Systems, v. 10, n. 5, p. 557-570, 2002.",
        "TUFTE, E. R. The Visual Display of Quantitative Information. 2. ed. Cheshire: Graphics Press, 2001.",
        "VASWANI, A.; SHAZEER, N.; PARMAR, N.; USZKOREIT, J.; JONES, L.; GOMEZ, A. N.; KAISER, L.; POLOSUKHIN, I. Attention is All You Need. Advances in Neural Information Processing Systems (NeurIPS), v. 30, p. 5998-6008, 2017.",
        "VOINEA, S. V.; MAMULEANU, M.; TEICA, R. V.; FLORESCU, L. M.; SELISTEANU, D.; GHEONEA, I. A. GPT-driven radiology report generation with fine-tuned Llama 3. Bioengineering, v. 11, n. 10, p. 1043, 2024.",
        "WIEGAND, I. C.; JUNGMANN, F.; HAN, T.; SIEPMANN, R.; KATHER, J. N. et al. Deidentifying medical documents with local, privacy-preserving large language models: the LLM-Anonymizer. NEJM AI, [s.l.], 2024.",
        "WOHLIN, C.; RUNESON, P.; HOST, M.; OHLSSON, M. C.; REGNELL, B.; WESSLEN, A. Experimentation in Software Engineering. Berlin: Springer, 2012.",
        "XIONG, G.; JIN, Q.; LU, Z.; ZHANG, A. Benchmarking retrieval-augmented generation for medicine. Findings of the Association for Computational Linguistics: ACL 2024, Bangkok, 2024, p. 6233-6251.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        set_pf(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0,
               space_after=Pt(6), first_line=Cm(0))
        run = p.add_run(r)
        run.font.name = "Arial"
        run.font.size = Pt(12)


def build_apendice_reprodutibilidade(doc):
    """Apêndice A: módulos de software e gold standards versionados."""
    add_h1(doc, "APÊNDICE A - MÓDULOS DE SOFTWARE PARA REPRODUTIBILIDADE")
    add_text(
        doc,
        "Este apêndice consolida os módulos Python publicados em "
        "conjunto com este trabalho, suficientes para reproduzir os "
        "resultados das Seções 6.1 a 6.5. O código-fonte está "
        "versionado e publicamente disponível em: "
        "https://github.com/Alencar-png/emotion-care-tcc. Os "
        "prompts de sistema dos agentes LLM constam no Apêndice B "
        "deste documento, na íntegra.",
        first_line=Cm(0))
    add_table(doc,
              ["Módulo", "Função", "Caminho relativo"],
              [["metrics_core.py",
                "Bootstrap CI, Fβ, span-level, McNemar, multiclass",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["regulatory_lookup.py",
                "Base curada NR-01/NR-17/COPSOQ/ISO 45003 para "
                "verificação de alucinação regulatória",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["reference_inventory_generator.py",
                "Gerador determinístico de inventário-referência "
                "para BERTScore/ROUGE do narrador",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["likert_rubrics.py",
                "Rubricas Likert pré-registradas + kappa de Cohen "
                "ponderado + exportador de amostras estratificadas",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["pii_gold_1000.py",
                "Gold do validador anti-PII: 1.025 casos com "
                "anotação span-level (445+/580-)",
                "emotion-care-tcc/refactor_v1_2/eval/golds/"],
               ["action_plan_gold_v3.py",
                "Gold do 5W2H: 80 cenários cegos (20 por nível NR-01) "
                "com rubrica explícita",
                "emotion-care-tcc/refactor_v1_2/eval/golds/"],
               ["eval_pii_v3.py",
                "Avaliação do anti-PII com Fβ=2, "
                "span-level e taxa de vazamento",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["eval_narrator_v3.py",
                "Avaliação do narrador com BERTScore, ROUGE, "
                "faithfulness e alucinação regulatória",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["eval_action_plan_v3.py",
                "Avaliação do 5W2H sobre gold cego com matriz 4x4 "
                "e diversidade",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["eval_copilot_v3.py",
                "Avaliação do copiloto sob framework RAGAS "
                "(faithfulness, answer relevancy, context relevancy)",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["eval_qualitative_v3.py",
                "Avaliação do analisador qualitativo com c_v, "
                "stability across seeds e baselines BERTopic/k-means/LDA",
                "emotion-care-tcc/refactor_v1_2/eval/"],
               ["pii_validator.py",
                "Implementação do validador anti-PII (regex + "
                "heurística + lista de prenomes brasileiros)",
                "emotion-care-tcc/emotion-care/services/ai/"],
               ["pgr_narrator.py",
                "Implementação do narrador do PGR (chamada LLM + "
                "re-prompting iterativo + validação de saída)",
                "emotion-care-tcc/emotion-care/services/ai/"],
               ["action_plan_generator.py",
                "Implementação do gerador 5W2H com prompt "
                "estruturado e validação anti-PII pós-geração",
                "emotion-care-tcc/emotion-care/services/ai/"],
               ["copilot_nr01.py",
                "Implementação do copiloto NR-01 com RAG sobre "
                "pgvector + geração condicionada",
                "emotion-care-tcc/emotion-care/services/ai/"],
               ["qualitative_analyzer.py",
                "Implementação do analisador qualitativo com "
                "clustering temático via LLM",
                "emotion-care-tcc/emotion-care/services/ai/"],
               ["anonymity_policy.py",
                "Política de k-anonymity (k=4 por dimensão, k=3 "
                "por setor) inegociável",
                "emotion-care-tcc/emotion-care/services/"]],
              "Quadro 2 - Módulos de software para reprodutibilidade",
              "Fonte: elaborado pelos autores (2026). Cada módulo é "
              "versionado em Git e contém docstrings descrevendo "
              "entradas, saídas e dependências.")


def _extract_system_prompt(path: Path) -> str:
    """Extrai a variável SYSTEM_PROMPT de um arquivo Python do projeto."""
    if not path.exists():
        return "(prompt indisponível)"
    src = path.read_text(encoding="utf-8")
    import re as _re
    m = _re.search(
        r'(?:SYSTEM_PROMPT|CLUSTER_SYSTEM_PROMPT)\s*=\s*"""(.+?)"""',
        src, _re.DOTALL,
    )
    return m.group(1).strip() if m else "(prompt não localizado)"


def build_apendice_prompts(doc):
    add_h1(doc, "APÊNDICE B - PROMPTS DE SISTEMA DOS AGENTES LLM")
    add_text(
        doc,
        "Este apêndice apresenta, na íntegra, os system prompts "
        "utilizados nos quatro agentes baseados em LLM (o validador "
        "anti-PII é determinístico e não utiliza prompt). A "
        "reprodutibilidade dos experimentos descritos no Capítulo 6 "
        "depende destes prompts em conjunto com os módulos listados "
        "no Apêndice A e versionados em "
        "https://github.com/Alencar-png/emotion-care-tcc.",
        first_line=Cm(0))

    backend = ROOT.parent / "emotion-care" / "nr1-backend" / "services" / "ai"
    agents = [
        ("B.1 Narrador do PGR", backend / "pgr_narrator.py"),
        ("B.2 Gerador de Plano de Ação 5W2H",
         backend / "action_plan_generator.py"),
        ("B.3 Copiloto NR-01 com RAG", backend / "copilot_nr01.py"),
        ("B.4 Analisador Qualitativo", backend / "qualitative_analyzer.py"),
    ]
    for title, path in agents:
        add_h2(doc, title)
        prompt_text = _extract_system_prompt(path)
        # Bloco de citação direta longa: fonte 10, espacamento simples,
        # recuo de 2 cm conforme ABNT NBR 10520.
        for paragraph in prompt_text.split("\n"):
            p = doc.add_paragraph()
            set_pf(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0,
                   space_after=Pt(0), first_line=Cm(0),
                   left_indent=Cm(2))
            r = p.add_run(paragraph if paragraph.strip() else " ")
            r.font.name = "Courier New"
            r.font.size = Pt(9)
        doc.add_paragraph()

    add_h2(doc, "B.5 Validador anti-PII")
    add_text(
        doc,
        "O validador anti-PII não utiliza prompt LLM. Trata-se de "
        "componente determinístico baseado em expressões regulares e "
        "heurísticas semânticas (lista de prenomes brasileiros mais "
        "padrões de título com prenome) implementado em "
        "emotion-care-tcc/emotion-care/services/ai/pii_validator.py. "
        "As oito categorias detectadas (e-mail, CPF, CNPJ, telefone, "
        "matrícula, cargo identificador, título com prenome e nome "
        "composto) e a lógica de redação estão documentadas no "
        "Capítulo 6 (Seção 6.1) e no código-fonte do repositório."
    )


# --------- conversão PDF ---------

def convert_to_pdf(docx_path: Path) -> Path | None:
    """Converte .docx em PDF via LibreOffice headless."""
    import shutil
    import subprocess
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        shutil.which("soffice"),
        shutil.which("libreoffice"),
    ]
    soffice = next((c for c in candidates if c and Path(c).exists()), None)
    if soffice is None:
        print("[!] LibreOffice não encontrado; PDF não gerado.")
        return None

    out_dir = docx_path.parent
    cmd = [soffice, "--headless", "--convert-to", "pdf",
           "--outdir", str(out_dir), str(docx_path)]
    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=180)
    except subprocess.CalledProcessError as e:
        msg = e.stderr.decode(errors='replace')[:500] if e.stderr else "(no stderr)"
        print(f"[!] LibreOffice falhou: {msg}")
        return None

    pdf_path = out_dir / (docx_path.stem + ".pdf")
    return pdf_path if pdf_path.exists() else None


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    set_margins(doc.sections[0])

    build_pretextuais(doc)
    build_introducao(doc)
    add_pagebreak(doc)
    build_objetivos(doc)
    add_pagebreak(doc)
    build_aspectos_teoricos(doc)
    add_pagebreak(doc)
    build_trabalhos_relacionados(doc)
    add_pagebreak(doc)
    build_metodologia(doc)
    add_pagebreak(doc)
    build_resultados(doc)
    add_pagebreak(doc)
    build_discussao(doc)
    add_pagebreak(doc)
    build_conclusoes(doc)
    add_pagebreak(doc)
    build_trabalhos_futuros(doc)
    add_pagebreak(doc)
    build_referencias(doc)
    add_pagebreak(doc)
    build_apendice_reprodutibilidade(doc)
    add_pagebreak(doc)
    build_apendice_prompts(doc)

    doc.save(OUT)
    size_kb = OUT.stat().st_size / 1024
    print(f"OK -> {OUT} ({size_kb:.1f} KB)")

    pdf = convert_to_pdf(OUT)
    if pdf:
        pdf_kb = pdf.stat().st_size / 1024
        print(f"PDF -> {pdf} ({pdf_kb:.1f} KB)")


if __name__ == "__main__":
    main()
